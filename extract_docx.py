import sys
import docx

# Set UTF-8 encoding for stdout
sys.stdout.reconfigure(encoding='utf-8')

# Extract from first document
doc1 = docx.Document(r'C:\Cursor\TayfaProject\EquipmentCounter\Data\test2\VOR_generated.docx')
print("=" * 80)
print("VOR_generated.docx")
print("=" * 80)
for p in doc1.paragraphs:
    print(p.text)

print("\n" + "=" * 80)
print("=" * 80)
print("\n")

# Extract from second document
doc2 = docx.Document(r'C:\Cursor\TayfaProject\EquipmentCounter\Data\test2\ВОР ЭОМ_29.docx')
print("=" * 80)
print("ВОР ЭОМ_29.docx")
print("=" * 80)
for p in doc2.paragraphs:
    print(p.text)
