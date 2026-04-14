import sys
import docx

# Set UTF-8 encoding for stdout
sys.stdout.reconfigure(encoding='utf-8')

def extract_document_content(file_path, doc_name):
    print("=" * 80)
    print(f"{doc_name}")
    print("=" * 80)

    doc = docx.Document(file_path)

    # Extract paragraphs
    print("\n--- ПАРАГРАФЫ ---")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"{i+1}. {p.text}")

    # Extract tables
    print("\n--- ТАБЛИЦЫ ---")
    for table_idx, table in enumerate(doc.tables):
        print(f"\nТаблица {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            cells_text = []
            for cell in row.cells:
                cells_text.append(cell.text.strip())
            print(f"  Строка {row_idx + 1}: {' | '.join(cells_text)}")

    print("\n" + "=" * 80 + "\n")

# Extract from first document
extract_document_content(
    r'C:\Cursor\TayfaProject\EquipmentCounter\Data\test2\VOR_generated.docx',
    'VOR_generated.docx'
)

# Extract from second document
extract_document_content(
    r'C:\Cursor\TayfaProject\EquipmentCounter\Data\test2\ВОР ЭОМ_29.docx',
    'ВОР ЭОМ_29.docx'
)
