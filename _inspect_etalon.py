"""Глубокая инспекция эталонного ВОР: шрифты, объединения, стили."""
import sys
from docx import Document
from docx.shared import Emu

sys.stdout.reconfigure(encoding='utf-8')

p = r'Data\ДБТ разделы для ИИ\03_ГПК_\3-я захватка\ВОР ЭО, Захватка 3_ГПК.docx'
d = Document(p)

print('=== SECTION ===')
sec = d.sections[0]
print(f'page: {sec.page_width.cm:.1f} x {sec.page_height.cm:.1f} cm')
print(f'margins: L={sec.left_margin.cm:.1f} R={sec.right_margin.cm:.1f} '
      f'T={sec.top_margin.cm:.1f} B={sec.bottom_margin.cm:.1f}')
print(f'orientation: {sec.orientation}')

print()
print('=== PARAGRAPHS (full) ===')
for i, par in enumerate(d.paragraphs):
    runs_info = []
    for r in par.runs:
        f = r.font
        runs_info.append(f'{(f.size.pt if f.size else "?")}pt b={f.bold} i={f.italic}')
    print(f'[{i}] align={par.alignment} style={par.style.name!r}')
    print(f'    text: {par.text[:200]!r}')
    if runs_info:
        print(f'    runs: {runs_info[:3]}')

print()
print('=== TABLE 0 ===')
tbl = d.tables[0]
print(f'rows={len(tbl.rows)} cols={len(tbl.columns)}')
print(f'alignment: {tbl.alignment}')
print(f'style: {tbl.style.name if tbl.style else None}')

# column widths
print()
print('Column widths:')
for ci, col in enumerate(tbl.columns):
    try:
        w = col.cells[0].width
        if w is not None:
            print(f'  col {ci}: {w.cm:.2f} cm')
        else:
            print(f'  col {ci}: None')
    except Exception as e:
        print(f'  col {ci}: err {e}')

# header row formatting
print()
print('Header row (R0) cell details:')
for ci, c in enumerate(tbl.rows[0].cells):
    txt = c.text.strip()
    paras = c.paragraphs
    runs = [r for p in paras for r in p.runs]
    bolds = [r.bold for r in runs[:3]]
    sizes = [(r.font.size.pt if r.font.size else None) for r in runs[:3]]
    print(f'  C{ci}: {txt!r:50s} bold={bolds} size={sizes}')

# data row sample
print()
print('Sample data rows (R3, R7=header-ish):')
for ri in [3, 7, 8]:
    if ri < len(tbl.rows):
        row = tbl.rows[ri]
        print(f'\nRow {ri}:')
        for ci, c in enumerate(row.cells):
            txt = c.text.strip()
            runs = [r for par in c.paragraphs for r in par.runs]
            bolds = list({r.bold for r in runs})
            sizes = list({(r.font.size.pt if r.font.size else None) for r in runs})
            print(f'  C{ci}: bold={bolds} size={sizes} | {txt[:60]!r}')

# Detect group/section rows (where col2 has text but cols 1,3,4 empty)
print()
print('=== SECTION-LIKE ROWS (group headers) ===')
for ri, row in enumerate(tbl.rows):
    cells = [c.text.strip() for c in row.cells]
    if cells[1] and not cells[0] and not cells[2] and not cells[3]:
        print(f'  R{ri}: {cells[1]!r}')

# Print row count by detecting non-empty rows
print()
print('=== LAST 10 ROWS ===')
for ri in range(max(0, len(tbl.rows) - 10), len(tbl.rows)):
    row = tbl.rows[ri]
    cells = [c.text.strip()[:40] for c in row.cells]
    print(f'  R{ri}: {cells}')
