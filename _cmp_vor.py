"""Compare two VOR DOCX files by per-item quantities.

Left:  VOR_02_PDF (3).docx        -- our generated VOR (S019/S020 output)
Right: ВОР ЭО, Захватка 3_ГПК.docx -- human-authored reference

Strategy:
  - For each docx, walk table(s), extract rows.
  - Identify columns: №, Наименование, Ед.изм., Количество (Объем работ).
  - Group rows by normalized name; sum quantities.
  - Produce side-by-side comparison.
"""
from __future__ import annotations

import re
from pathlib import Path
from docx import Document

LEFT = Path(r"C:\Cursor\TayfaProject\EquipmentCounter\Data\ДБТ разделы для ИИ\03_ГПК_\3-я захватка\VOR_02_PDF_3.docx")
RIGHT = Path(r"C:\Cursor\TayfaProject\EquipmentCounter\Data\ДБТ разделы для ИИ\03_ГПК_\3-я захватка\ВОР ЭО, Захватка 3_ГПК.docx")
OUT = Path(r"C:\Cursor\TayfaProject\EquipmentCounter\_cmp_vor_out.txt")


def parse_float(s: str) -> float | None:
    if not s:
        return None
    s = s.strip().replace(",", ".").replace("\xa0", "").replace(" ", "")
    if not s:
        return None
    try:
        return float(s)
    except ValueError:
        # try extract first number
        m = re.search(r"[+-]?\d+(?:\.\d+)?", s)
        if m:
            try:
                return float(m.group(0))
            except ValueError:
                return None
        return None


def cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def extract_rows(docx_path: Path) -> list[dict]:
    """Walk all tables, return list of {name, unit, qty, raw} for data rows."""
    d = Document(str(docx_path))
    rows_out: list[dict] = []
    for ti, t in enumerate(d.tables):
        # Try to detect column layout from first row
        if not t.rows:
            continue
        header_cells = [cell_text(c) for c in t.rows[0].cells]
        # Find name/unit/qty column indices
        name_i = unit_i = qty_i = None
        for i, h in enumerate(header_cells):
            hl = h.lower()
            if name_i is None and ("наимен" in hl or "вид работ" in hl):
                name_i = i
            if unit_i is None and ("ед" in hl and "изм" in hl):
                unit_i = i
            if qty_i is None and ("объем работ" in hl or "количество" in hl or "кол-во" in hl):
                qty_i = i
        # Defaults if header detection failed: positional 2/3/4 (after № at idx 0/1)
        if name_i is None:
            name_i = 1
        if unit_i is None:
            unit_i = 2
        if qty_i is None:
            qty_i = 3

        for ri, row in enumerate(t.rows):
            if ri == 0:
                continue
            cells = [cell_text(c) for c in row.cells]
            # Some rows may have fewer cells (merged); pad
            while len(cells) <= max(name_i, unit_i, qty_i):
                cells.append("")
            name = cells[name_i].strip()
            unit = cells[unit_i].strip()
            qty_raw = cells[qty_i].strip()
            # Skip numeric-header rows like "1 2 3 4 5 6 7"
            if all(c.strip().isdigit() and len(c.strip()) <= 2 for c in cells if c.strip()):
                continue
            if not name:
                continue
            qty = parse_float(qty_raw)
            rows_out.append({
                "table": ti,
                "row": ri,
                "name": name,
                "unit": unit,
                "qty": qty,
                "qty_raw": qty_raw,
            })
    return rows_out


def norm_name(name: str) -> str:
    """Normalize for matching: lowercase, strip extras, collapse whitespace."""
    s = name.lower()
    s = re.sub(r"\s+", " ", s)
    s = s.strip(" .,:-—")
    return s


def short_key(name: str) -> str:
    """Even shorter key: first 3-4 significant words for fuzzy matching."""
    n = norm_name(name)
    # strip common qualifiers
    n = re.sub(r"\bна высоте.+", "", n)
    n = re.sub(r"\bвыс\.\s*\d+.+", "", n)
    n = re.sub(r"\bдиам\.\s*\d+.+", "", n)
    n = re.sub(r"\bсеч\.\s*\d+.+", "", n)
    n = re.sub(r"\b\d+x\d+.+", "", n)
    words = n.split()[:5]
    return " ".join(words).strip()


def aggregate(rows: list[dict]) -> dict[str, dict]:
    """Aggregate by short_key -> {name, unit, qty_total, count, samples}."""
    agg: dict[str, dict] = {}
    for r in rows:
        k = short_key(r["name"])
        if not k:
            continue
        if k not in agg:
            agg[k] = {
                "key": k,
                "name_sample": r["name"],
                "unit": r["unit"],
                "qty_total": 0.0,
                "count": 0,
                "names_seen": set(),
            }
        agg[k]["names_seen"].add(r["name"])
        agg[k]["count"] += 1
        if r["qty"] is not None:
            agg[k]["qty_total"] += r["qty"]
    return agg


def main():
    out_lines = []
    out_lines.append("=" * 100)
    out_lines.append("COMPARISON: VOR_02_PDF_3.docx  vs  ВОР ЭО, Захватка 3_ГПК.docx")
    out_lines.append("=" * 100)

    left_rows = extract_rows(LEFT)
    right_rows = extract_rows(RIGHT)
    out_lines.append(f"LEFT  (ours):     {len(left_rows)} data rows")
    out_lines.append(f"RIGHT (etalon):   {len(right_rows)} data rows")
    out_lines.append("")

    L = aggregate(left_rows)
    R = aggregate(right_rows)
    out_lines.append(f"LEFT  aggregated keys: {len(L)}")
    out_lines.append(f"RIGHT aggregated keys: {len(R)}")
    out_lines.append("")

    # Union of keys for side-by-side
    all_keys = sorted(set(L.keys()) | set(R.keys()))

    out_lines.append("-" * 100)
    out_lines.append(f"{'KEY':<55} {'LEFT_qty':>10} {'RIGHT_qty':>10} {'DIFF':>10} {'UNIT_L':<8} {'UNIT_R':<8}")
    out_lines.append("-" * 100)

    matched_both = 0
    only_left = 0
    only_right = 0
    for k in all_keys:
        l = L.get(k)
        r = R.get(k)
        lq = f"{l['qty_total']:.1f}" if l else "—"
        rq = f"{r['qty_total']:.1f}" if r else "—"
        if l and r:
            diff = l["qty_total"] - r["qty_total"]
            diff_str = f"{diff:+.1f}"
            matched_both += 1
        elif l and not r:
            diff_str = "ONLY-L"
            only_left += 1
        else:
            diff_str = "ONLY-R"
            only_right += 1
        ul = (l["unit"] if l else "")[:8]
        ur = (r["unit"] if r else "")[:8]
        out_lines.append(f"{k[:55]:<55} {lq:>10} {rq:>10} {diff_str:>10} {ul:<8} {ur:<8}")

    out_lines.append("-" * 100)
    out_lines.append(f"matched_in_both = {matched_both}")
    out_lines.append(f"only_in_left    = {only_left}")
    out_lines.append(f"only_in_right   = {only_right}")
    out_lines.append("")

    # Also dump raw rows for cross-check
    out_lines.append("=" * 100)
    out_lines.append("LEFT RAW (ours)")
    out_lines.append("=" * 100)
    for r in left_rows[:200]:
        out_lines.append(f"  {r['name'][:80]:<80} | {r['unit']:<8} | {r['qty_raw']!r}")
    if len(left_rows) > 200:
        out_lines.append(f"... ({len(left_rows)-200} more)")
    out_lines.append("")
    out_lines.append("=" * 100)
    out_lines.append("RIGHT RAW (etalon)")
    out_lines.append("=" * 100)
    for r in right_rows[:200]:
        out_lines.append(f"  {r['name'][:80]:<80} | {r['unit']:<8} | {r['qty_raw']!r}")
    if len(right_rows) > 200:
        out_lines.append(f"... ({len(right_rows)-200} more)")

    OUT.write_text("\n".join(out_lines), encoding="utf-8")


if __name__ == "__main__":
    main()
