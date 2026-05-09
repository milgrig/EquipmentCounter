#!/usr/bin/env python3
"""
vor_height_fix.py
=================

Постпроцессор для исправления распределения светильников и кабелей
по категориям высоты в сгенерированном ВОР_ЭО.docx.

Проблема
--------
В сгенерированном ВОР категории высоты заполнены по правилу:
    elevation_to_height(отметка_пола_этажа)
Это приводит к тому, что светильник с пола +4.200 попадает в «до 5 метров»,
хотя реально монтажная высота = отметка потолка этажа (пол следующего) =
+7.800 → «от 5 до 13 метров».

Решение
-------
1) Читаем сгенерированный ВОР и находим все строки вида
   "Монтаж светильников ... на высоте {категория}:"  и
   "Прокладка кабеля в лотке на высоте {категория}:"  и пр.
2) Читаем исходные данные: список имён файлов захватки и, опционально,
   equipment_report.json с разбивкой по отметкам.
3) Вычисляем правильное распределение через height_mapping.py.
4) Пересчитываем количества в строках ВОР.
5) Сохраняем исправленный DOCX.

Этот скрипт НЕ модифицирует vor_generator.py или equipment_counter.py —
он работает как внешний постпроцессор поверх их вывода.

Использование
-------------
    python vor_height_fix.py <captures_dir> <input_docx> [<output_docx>]

Пример:
    python vor_height_fix.py \
      "Data/ДБТ разделы для ИИ/03_ГПК_/3-я захватка" \
      "Data/ДБТ разделы для ИИ/03_ГПК_/3-я захватка/ВОР_ЭО.docx" \
      "Data/ДБТ разделы для ИИ/03_ГПК_/3-я захватка/ВОР_ЭО_height_fix.docx"
"""
from __future__ import annotations

import argparse
import io
import json
import os
import re
import sys
from pathlib import Path
from typing import Iterable

from height_mapping import (
    build_floor_to_ceiling_map,
    collect_floors_from_filenames,
    extract_elevations_from_filename,
    floor_to_height_category,
    height_to_category,
    remap_counts_by_floor,
)

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Категории, встречающиеся в ВОР
# ---------------------------------------------------------------------------

CATEGORIES_4 = [
    "до 5 метров",
    "от 5 до 13 метров",
    "от 13 до 20 метров",
    "от 20 до 35 метров",
]


# Паттерн категории высоты в строке: "до 5 метров" / "от 5 до 13 метров"
CATEGORY_RE = re.compile(
    r"(до\s*5\s*(?:метров|м)|"
    r"от\s*5\s*до\s*13\s*(?:метров|м)|"
    r"от\s*13\s*до\s*20\s*(?:метров|м)|"
    r"от\s*20\s*до\s*35\s*(?:метров|м))",
    re.IGNORECASE,
)


# ---------------------------------------------------------------------------
# Сбор данных: отметки этажей
# ---------------------------------------------------------------------------


def collect_capture_floors(captures_dir: Path) -> list[float]:
    """
    Сканирует папку захватки (DWG/DXF) и возвращает все отметки этажей
    (по именам файлов планов освещения / привязки / кабеленесущих).
    """
    if not captures_dir.exists():
        return []
    filenames: list[str] = []
    for ext in ("*.dwg", "*.dxf"):
        for f in captures_dir.rglob(ext):
            # Интересуют только планы с привязкой к отметке
            name = f.name.lower()
            if ("отм" in name and ("план" in name or "освещен" in name or
                                    "привязк" in name or "кабеленесущ" in name)):
                filenames.append(f.name)
    return collect_floors_from_filenames(filenames)


def collect_per_floor_luminaires(
    equipment_report_path: Path,
) -> dict[float, int]:
    """
    Читает equipment_report.json и собирает по каждому файлу «план привязки»
    количество светильников.  Возвращает {отметка_пола: суммарное_кол-во}.

    Фильтрует только файлы с «привязки светильников» в имени — они содержат
    канонические названия (Светильник SLICK.PRS LED ...).
    """
    if not equipment_report_path.exists():
        return {}
    try:
        data = json.loads(equipment_report_path.read_text(encoding="utf-8"))
    except Exception as e:
        print(f"[warn] не удалось прочитать {equipment_report_path.name}: {e}")
        return {}

    per_floor: dict[float, int] = {}
    for fname, info in data.get("files", {}).items():
        name_low = fname.lower()
        if "привязк" not in name_low:
            continue
        elevs = extract_elevations_from_filename(fname)
        if not elevs:
            continue
        floor = elevs[0]  # для файлов привязки — одна отметка
        count = 0
        for it in info.get("equipment", []):
            nm = it.get("name", "")
            if ("светильник" in nm.lower() or
                "пиктограмма" in nm.lower() or
                "указател" in nm.lower()):
                count += int(it.get("total", 0))
        per_floor[floor] = per_floor.get(floor, 0) + count
    return per_floor


# ---------------------------------------------------------------------------
# Работа с DOCX
# ---------------------------------------------------------------------------


def load_docx_tables(path: Path):
    """Загружает DOCX и возвращает Document + список таблиц."""
    try:
        from docx import Document
    except ImportError:
        print("[error] python-docx не установлен: pip install python-docx")
        sys.exit(2)
    doc = Document(str(path))
    return doc, list(doc.tables)


def find_main_vor_table(tables: list) -> object | None:
    """Находит главную ВОР-таблицу — самую длинную (с 5+ колонками)."""
    best = None
    best_rows = 0
    for t in tables:
        if not t.rows:
            continue
        n_cols = len(t.rows[0].cells)
        if n_cols < 5:
            continue
        if len(t.rows) > best_rows:
            best = t
            best_rows = len(t.rows)
    return best


def detect_category(text: str) -> str | None:
    """Возвращает каноническую метку категории из текста."""
    m = CATEGORY_RE.search(text)
    if not m:
        return None
    raw = m.group(1).lower().replace("\u00a0", " ")
    raw = re.sub(r"\s+", " ", raw)
    # нормализуем 'м' → 'метров'
    for cat in CATEGORIES_4:
        cat_norm = cat.lower()
        # паттерн без окончания (метров/м)
        cat_prefix = cat_norm.rsplit(" ", 1)[0]
        if raw.startswith(cat_prefix):
            return cat
    return None


def classify_work_row(name: str) -> str | None:
    """
    Определяет тип работы в строке ВОР.  Возвращает:
      - 'luminaire_stud'    — шпильки к перекрытию
      - 'luminaire_wall'    — настенный
      - 'luminaire_anchor'  — анкерный
      - 'luminaire_pendant' — подвесной
      - 'cable_tray'        — в лотке
      - 'cable_gofra'       — в гофре
      - 'tray_mount'        — сам лоток
      - 'gofra_mount'       — монтаж гофры
      - None — не наш случай
    """
    nl = name.lower()
    if "монтаж светильник" in nl and "шпильк" in nl:
        return "luminaire_stud"
    if "монтаж настенного" in nl and "светильник" in nl:
        return "luminaire_wall"
    if "монтаж светильник" in nl and "анкер" in nl:
        return "luminaire_anchor"
    if "монтаж светильник" in nl and "подвесн" in nl:
        return "luminaire_pendant"
    if "прокладка кабел" in nl and "лотк" in nl:
        return "cable_tray"
    if "прокладка кабел" in nl and "гофр" in nl:
        return "cable_gofra"
    if "лоток" in nl and "высот" in nl and "лотка" in nl:
        return "tray_mount"
    if "монтаж гофрирован" in nl and "высот" in nl:
        return "gofra_mount"
    return None


# ---------------------------------------------------------------------------
# Применение пересчёта
# ---------------------------------------------------------------------------


def compute_correct_luminaire_distribution(
    per_floor_counts: dict[float, int],
    all_floors: list[float],
) -> dict[str, int]:
    """Пересчитывает counts по правильной монтажной высоте."""
    mapping = build_floor_to_ceiling_map(all_floors)
    return remap_counts_by_floor(per_floor_counts, mapping)


def rebalance_by_etalon_proportion(
    current_by_cat: dict[str, int],
    target_totals_by_cat: dict[str, int],
) -> dict[str, int]:
    """
    Когда у нас общая сумма одинакова, но распределение перекошено,
    пересчитать можно, перекинув количество из «лишних» категорий в
    «недобранные» согласно пропорциям целевого распределения.
    Эта функция не используется в текущей версии, оставлена для future work.
    """
    total_current = sum(current_by_cat.values())
    total_target = sum(target_totals_by_cat.values())
    if total_target == 0:
        return current_by_cat
    scale = total_current / total_target
    return {cat: round(v * scale) for cat, v in target_totals_by_cat.items()}


# ---------------------------------------------------------------------------
# Основная функция
# ---------------------------------------------------------------------------


def process_vor(
    capture_dir: Path,
    input_docx: Path,
    output_docx: Path,
    verbose: bool = True,
) -> dict:
    """
    Основной метод: читает ВОР, пересчитывает категории высот, сохраняет
    исправленный DOCX.  Возвращает словарь с метриками.
    """
    from docx.shared import Pt

    print(f"\n{'='*70}")
    print(f"VOR Height Fix Postprocessor")
    print(f"{'='*70}")
    print(f"Capture:  {capture_dir}")
    print(f"Input:    {input_docx.name}")
    print(f"Output:   {output_docx.name}")

    # 1) Сбор отметок
    floors = collect_capture_floors(capture_dir)
    if not floors:
        print("[warn] Не удалось извлечь отметки из имён файлов.  Используем ГПК-3 по умолчанию.")
        floors = [0.0, 4.2, 7.8, 9.0, 13.8, 18.6, 23.4, 28.2]
    print(f"\nОтметки этажей в захватке: {floors}")

    mapping = build_floor_to_ceiling_map(floors)
    print(f"\nКарта floor → ceiling:")
    for f in sorted(mapping):
        cat = height_to_category(mapping[f])
        print(f"  пол +{f:6.3f}  →  потолок +{mapping[f]:6.3f}  →  категория «{cat}»")

    # 2) Сбор per-floor светильников из equipment_report.json
    eq_report = capture_dir / "equipment_report.json"
    per_floor_lum = collect_per_floor_luminaires(eq_report)
    print(f"\nСветильники по отметкам пола (из equipment_report.json):")
    for f in sorted(per_floor_lum):
        print(f"  +{f:6.3f} → {per_floor_lum[f]:4d} шт")
    total_lum_by_floor = sum(per_floor_lum.values())
    print(f"  ИТОГО: {total_lum_by_floor} шт")

    corrected = remap_counts_by_floor(per_floor_lum, mapping)
    print(f"\nПравильное распределение по категориям:")
    for cat in CATEGORIES_4:
        print(f"  {cat:25}: {corrected.get(cat, 0):4d} шт")

    # 3) Открываем ВОР
    doc, tables = load_docx_tables(input_docx)
    main_table = find_main_vor_table(tables)
    if main_table is None:
        print("[error] Не найдена главная ВОР-таблица")
        return {}

    # 4) Пройти все строки, классифицировать, пересчитать для светильников
    updates: list[dict] = []  # что изменено
    # Структура ВОР (столбцы):
    #   0: № п/п
    #   1: Наименование вида работ
    #   2: Ед. изм.
    #   3: Объём работ  ← НУЖНОЕ КОЛИЧЕСТВО
    #   4: Формула
    #   5: Ссылка
    #   6: Доп. инфо
    NAME_COL = 1
    QTY_COL = 3
    # Соберём по типу работы текущие counts
    current_by_kind_cat: dict[tuple[str, str], tuple[int, object]] = {}
    # key=(kind, category), value=(current_qty, row_ref)
    for row in main_table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) <= max(NAME_COL, QTY_COL):
            continue
        name = cells[NAME_COL]
        kind = classify_work_row(name)
        if kind is None:
            continue
        cat = detect_category(name)
        if cat is None:
            continue
        qty_text = cells[QTY_COL].replace(" ", "").replace("\u00a0", "")
        try:
            qty = int(float(qty_text.replace(",", ".")))
        except ValueError:
            qty = 0
        current_by_kind_cat[(kind, cat)] = (qty, row)

    # 5) Пересчёт: для каждого типа светильника пропорционально исправляем
    # distribution.  Делаем ПО СУММЕ: сохраняем общую сумму для каждого kind,
    # но перераспределяем по корректным категориям в пропорции corrected.
    kinds_to_rebalance = [
        "luminaire_stud",
        "luminaire_wall",
        "luminaire_anchor",
        "luminaire_pendant",
    ]

    # Для кабелей и лотков та же стратегия: пересчёт веса по количеству
    # светильников (пропорция высот) — это приблизительно, но лучше старого
    cable_kinds = ["cable_tray", "cable_gofra", "tray_mount", "gofra_mount"]

    def rebalance_distribution(
        kind: str, current: dict[str, int], target_weights: dict[str, int]
    ) -> dict[str, int]:
        """Перераспределить total по target_weights-пропорции."""
        total = sum(current.values())
        total_w = sum(target_weights.values())
        if total == 0 or total_w == 0:
            return current
        new = {}
        acc = 0
        cats = CATEGORIES_4
        for i, cat in enumerate(cats):
            w = target_weights.get(cat, 0)
            if i == len(cats) - 1:
                new[cat] = total - acc  # остаток, чтобы сумма точно совпала
            else:
                v = round(total * w / total_w)
                new[cat] = v
                acc += v
        return new

    all_kinds = kinds_to_rebalance + cable_kinds
    for kind in all_kinds:
        current = {
            cat: current_by_kind_cat.get((kind, cat), (0, None))[0]
            for cat in CATEGORIES_4
        }
        if sum(current.values()) == 0:
            continue
        new_dist = rebalance_distribution(kind, current, corrected)

        # Применить: обновить row по каждой категории
        for cat in CATEGORIES_4:
            pair = current_by_kind_cat.get((kind, cat))
            if pair is None:
                continue
            old_qty, row = pair
            new_qty = new_dist.get(cat, 0)
            if new_qty == old_qty:
                continue
            # Обновляем строго колонку QTY_COL
            c = row.cells[QTY_COL]
            try:
                if c.paragraphs:
                    para = c.paragraphs[0]
                    if para.runs:
                        para.runs[0].text = str(new_qty)
                        for r in para.runs[1:]:
                            r.text = ""
                    else:
                        para.add_run(str(new_qty))
                else:
                    c.text = str(new_qty)
            except Exception:
                pass
            updates.append({
                "kind": kind,
                "category": cat,
                "old": old_qty,
                "new": new_qty,
            })

    # 6) Сохраняем
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))

    print(f"\nИзменено строк: {len(updates)}")
    if verbose:
        print("\nДетализация изменений:")
        for u in updates:
            print(f"  [{u['kind']:20}] «{u['category']:18}»: {u['old']:>5} → {u['new']:>5}")

    return {
        "floors": floors,
        "per_floor_luminaires": per_floor_lum,
        "corrected_distribution": corrected,
        "updates": updates,
        "output": str(output_docx),
    }


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main() -> None:
    p = argparse.ArgumentParser(description="Постпроцессор ВОР: исправление высот")
    p.add_argument("capture_dir", type=Path, help="Папка захватки (с DWG и equipment_report.json)")
    p.add_argument("input_docx", type=Path, help="Входной ВОР_ЭО.docx")
    p.add_argument("output_docx", type=Path, nargs="?", default=None,
                   help="Выходной DOCX (по умолчанию рядом с суффиксом _height_fix)")
    args = p.parse_args()

    if args.output_docx is None:
        stem = args.input_docx.stem
        args.output_docx = args.input_docx.with_name(f"{stem}_height_fix.docx")

    process_vor(args.capture_dir, args.input_docx, args.output_docx)


if __name__ == "__main__":
    main()
