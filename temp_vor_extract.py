import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
import os, glob

# Read all VOR files
vor_files = glob.glob('Data/test/ВОР_ЭО*.docx') + glob.glob('Data/test/ВОР ЭОМ*.docx')
print(f"Found {len(vor_files)} VOR files\n")

for fpath in sorted(vor_files):
    fname = os.path.basename(fpath)
    print(f"\n{'='*80}")
    print(f"FILE: {fname}")
    print(f"{'='*80}")
    
    doc = Document(fpath)
    
    for ti, table in enumerate(doc.tables):
        print(f"\n--- Table {ti+1} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
            print(f"  Row {ri:3d}: {cells}")
    
    print(f"\n--- Paragraphs ---")
    for p in doc.paragraphs:
        if p.text.strip():
            print(f"  {p.text.strip()}")
