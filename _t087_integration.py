"""T087 integration check — render DOCX from the real T066 items cache.

Loads `.tayfa/common/discussions/T066_items_cache.json` (per-PDF item
lists produced by T073), simulates the path
`_aggregate_equipment` → `regroup_aggregate_by_elevation`
→ `render_vor_docx`, and dumps an `_t087_integration_sample.docx`.

Then re-opens the result and prints all bold rows in the table so we can
visually verify that all four height sub-headers appear in a real run.
"""
from __future__ import annotations

import io
import json
from collections import defaultdict
from pathlib import Path

from docx import Document

from vor_docx_renderer import render_vor_docx
from vor_elevation_grouper import regroup_aggregate_by_elevation, extract_elevation


CACHE = Path(".tayfa/common/discussions/T066_items_cache.json")
raw = json.loads(CACHE.read_text(encoding="utf-8"))


def _simulate_aggregate(per_pdf: dict[str, list[dict]]) -> list[dict]:
    """Minimal stand-in for web_app._aggregate_equipment.

    Groups raw items by (name, unit) summing total; builds per_file map.
    Drops zero-total rows.
    """
    buckets: dict[tuple, dict] = {}
    for pdf_name, items in per_pdf.items():
        for it in items:
            name = str(it.get("name") or "").strip()
            unit = str(it.get("unit") or "шт").strip()
            try:
                total = float(it.get("total") or 0)
            except (TypeError, ValueError):
                total = 0.0
            if total <= 0:
                continue
            key = (name, unit)
            b = buckets.setdefault(key, {
                "name": name,
                "unit": unit,
                "total": 0.0,
                "per_file": defaultdict(float),
                "drawing_refs_list": [],
            })
            b["total"] += total
            b["per_file"][pdf_name] += total
            if pdf_name not in b["drawing_refs_list"]:
                b["drawing_refs_list"].append(pdf_name)

    out: list[dict] = []
    for b in buckets.values():
        formula = "+".join(
            f"{int(v) if abs(v - round(v)) < 1e-6 else round(v, 1)}"
            for _, v in b["per_file"].items()
        )
        out.append({
            "name": b["name"],
            "unit": b["unit"],
            "total": b["total"],
            "formula": formula,
            "drawing_refs": ", ".join(b["drawing_refs_list"]),
            "per_file": dict(b["per_file"]),
        })
    return out


aggregated = _simulate_aggregate(raw)
print(f"Aggregated rows: {len(aggregated)}")

aggregated = regroup_aggregate_by_elevation(aggregated, labeled=False)
print(f"After elevation regroup: {len(aggregated)} rows")

bytes_ = render_vor_docx(
    aggregated,
    rel_folder="ЭО, Захватка 3_ГПК",
    issue_date="14.05.2026",
)
print(f"render_vor_docx returned {len(bytes_)} bytes")

doc = Document(io.BytesIO(bytes_))
tbl = doc.tables[0]
bold_rows: list[str] = []
for r in tbl.rows:
    cell = r.cells[1]
    is_bold = False
    for p in cell.paragraphs:
        for run in p.runs:
            if run.bold and run.text.strip():
                is_bold = True
                break
        if is_bold:
            break
    if is_bold:
        bold_rows.append(cell.text.strip())

print("\n=== All bold (section / sub-header) rows ===")
for b in bold_rows:
    print(f"  {b}")

required = [
    "Монтаж светильников на шпильках к перекрытию на высоте до 5 метров",
    "Монтаж светильников на шпильках к перекрытию на высоте от 5 до 13 метров",
    "Монтаж светильников на шпильках к перекрытию на высоте от 13 до 20 метров",
    "Монтаж светильников на шпильках к перекрытию на высоте от 20 до 35 метров",
]
missing = [h for h in required if h not in bold_rows]
present = [h for h in required if h in bold_rows]
print(f"\nHeight sub-headers present: {len(present)}/{len(required)}")
if missing:
    print(f"Missing (no luminaires in those buckets in the cache): {missing}")

out = Path("_t087_integration_sample.docx")
out.write_bytes(bytes_)
print(f"\nWrote integration sample: {out.resolve()}")
