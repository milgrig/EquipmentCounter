#!/usr/bin/env python3
"""
Generate detailed test report for ГПК3 3rd capture VOR comparison.
"""

import sys
from pathlib import Path
from datetime import datetime

sys.stdout.reconfigure(encoding="utf-8")

try:
    from docx import Document
    import openpyxl
except ImportError:
    sys.exit("Dependencies required: pip install python-docx openpyxl")

# Paths
BASE_DIR = Path(__file__).parent / "Data" / "ДБТ разделы для ИИ" / "03_ГПК_" / "3-я захватка"
ETALON_FILE = BASE_DIR / "ВОР ЭО, Захватка 3_ГПК.docx"
GENERATED_FILE = BASE_DIR / "VOR_generated_test.docx"
COMPARISON_FILE = BASE_DIR / "VOR_COMPARISON.xlsx"
REPORT_FILE = BASE_DIR / "TEST_REPORT.md"

def count_docx_rows(docx_path):
    """Count rows in DOCX table."""
    doc = Document(str(docx_path))
    if not doc.tables:
        return 0
    return len(doc.tables[0].rows)

def read_comparison_summary(xlsx_path):
    """Read summary from comparison Excel."""
    wb = openpyxl.load_workbook(str(xlsx_path), data_only=True)

    summary = {}

    # Read "Итого по разделам" sheet
    if "Итого по разделам" in wb.sheetnames:
        ws = wb["Итого по разделам"]
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row or not any(row):
                continue
            section = row[0]
            if section and section != "ИТОГО":
                summary[section] = {
                    'total_ref': row[1] if len(row) > 1 else 0,
                    'matched': row[2] if len(row) > 2 else 0,
                    'diff_qty': row[3] if len(row) > 3 else 0,
                    'missing': row[4] if len(row) > 4 else 0,
                    'extra': row[5] if len(row) > 5 else 0,
                }
            elif section == "ИТОГО":
                summary['_total'] = {
                    'total_ref': row[1] if len(row) > 1 else 0,
                    'matched': row[2] if len(row) > 2 else 0,
                    'diff_qty': row[3] if len(row) > 3 else 0,
                    'missing': row[4] if len(row) > 4 else 0,
                    'extra': row[5] if len(row) > 5 else 0,
                }

    wb.close()
    return summary

print("Генерация отчета о тестировании...")

# Collect data
etalon_rows = count_docx_rows(ETALON_FILE) if ETALON_FILE.exists() else 0
generated_rows = count_docx_rows(GENERATED_FILE) if GENERATED_FILE.exists() else 0
summary = read_comparison_summary(COMPARISON_FILE) if COMPARISON_FILE.exists() else {}

# Generate report
report = []
report.append("# ОТЧЕТ О ТЕСТИРОВАНИИ СИСТЕМЫ ГЕНЕРАЦИИ ВОР")
report.append("")
report.append(f"**Дата тестирования:** {datetime.now().strftime('%d.%m.%Y %H:%M')}")
report.append("")
report.append("## 1. Тестируемая система")
report.append("")
report.append("**Система:** EquipmentCounter - автоматическое составление ВОР из PDF-чертежей")
report.append("")
report.append("**Версия пайплайна:** run_web_v2 (с поддержкой вертикальных стояков)")
report.append("")
report.append("**Тестовый набор данных:**")
report.append(f"- Папка: `Data/ДБТ разделы для ИИ/03_ГПК_/3-я захватка/02_PDF`")
report.append(f"- PDF файлов: 31")
report.append(f"- Эталонный ВОР: `ВОР ЭО, Захватка 3_ГПК.docx`")
report.append("")
report.append("---")
report.append("")

report.append("## 2. Результаты обработки")
report.append("")
report.append(f"### 2.1 Статистика обработки PDF")
report.append("")
report.append("| Файл | Элементов найдено |")
report.append("|------|-------------------|")
report.append("| 003 - Схемы ЩО, ЩАО-ЩО-3.pdf | 3 |")
report.append("| 004 - Схемы ЩО, ЩАО-ЩАО-3.pdf | 1 |")
report.append("| 005-Планы освещения-отм. 0.000.pdf | 23 |")
report.append("| 006-Планы освещения-отм. +4.200.pdf | 27 |")
report.append("| 007-Планы освещения-отм. +7.800 +9.000.pdf | 19 |")
report.append("| 008-Планы освещения-отм. +13.800.pdf | 12 |")
report.append("| 009-Планы освещения-отм. +18.600.pdf | 17 |")
report.append("| 010-Планы освещения-отм. +23.400.pdf | 18 |")
report.append("| 011-Планы освещения-отм. +28.200.pdf | 19 |")
report.append("| 012-018 Планы привязки | 12-9 каждый |")
report.append("| 019-025 Лотки | 0 (не распознаются) |")
report.append("")
report.append(f"**Время обработки:** ~553 секунд (~9 минут)")
report.append("")
report.append(f"**Результат:** 26 уникальных позиций в ВОР")
report.append("")

report.append("### 2.2 Сравнение с эталоном")
report.append("")
report.append(f"| Параметр | Эталон | Сгенерировано | % от эталона |")
report.append(f"|----------|--------|---------------|--------------|")
report.append(f"| Строк всего | {etalon_rows} | {generated_rows} | {generated_rows/etalon_rows*100:.1f}% |")

if '_total' in summary:
    total = summary['_total']
    ref_items = total.get('total_ref', 0)
    matched = total.get('matched', 0)
    diff_qty = total.get('diff_qty', 0)
    missing = total.get('missing', 0)
    extra = total.get('extra', 0)

    report.append(f"| Позиций в эталоне | {ref_items} | - | - |")
    report.append(f"| Точные совпадения | - | {matched} | {matched/ref_items*100 if ref_items else 0:.1f}% |")
    report.append(f"| Расхождения в кол-ве | - | {diff_qty} | {diff_qty/ref_items*100 if ref_items else 0:.1f}% |")
    report.append(f"| Отсутствующие позиции | - | {missing} | {missing/ref_items*100 if ref_items else 0:.1f}% |")
    report.append(f"| Лишние позиции | - | {extra} | - |")
report.append("")

report.append("### 2.3 Детализация по разделам")
report.append("")
report.append("| Раздел | Эталон | Совпало | Расх-ние | Отсутствует | Лишние |")
report.append("|--------|--------|---------|----------|-------------|--------|")

for section, data in summary.items():
    if section.startswith('_'):
        continue
    ref = data.get('total_ref', 0)
    matched = data.get('matched', 0)
    diff_qty = data.get('diff_qty', 0)
    missing = data.get('missing', 0)
    extra = data.get('extra', 0)
    report.append(f"| {section} | {ref} | {matched} | {diff_qty} | {missing} | {extra} |")

report.append("")
report.append("---")
report.append("")

report.append("## 3. Анализ результатов")
report.append("")
report.append("### 3.1 ✅ Что работает")
report.append("")
report.append("1. **Планы освещения** — успешно распознаются")
report.append("   - Светильники различных типов идентифицируются")
report.append("   - Пиктограммы \"ВЫХОД\" и световые указатели находятся")
report.append("   - Агрегация по строительным отметкам работает")
report.append("")
report.append("2. **Планы привязки** — частично работают")
report.append("   - Находятся некоторые элементы из легенды")
report.append("")
report.append("3. **Схемы щитов** — минимальное распознавание")
report.append("   - Некоторые кабели извлекаются (3-4 элемента)")
report.append("")
report.append("4. **Вертикальные стояки** — добавляются корректно")
report.append("   - Модуль `vor_height_injector` отработал без ошибок")
report.append("")

report.append("### 3.2 ❌ Что не работает")
report.append("")
report.append("1. **Щитовое оборудование (100% отсутствует)**")
report.append("   - Не распознаются щиты: ЩО-3, ЩАО-3, ЦСАО3")
report.append("   - Не находится подключение жил кабелей")
report.append("   - Файлы схем (003, 004, 004.1) обрабатываются плохо")
report.append("")
report.append("2. **Установочные изделия (100% отсутствует)**")
report.append("   - Не распознаются розетки")
report.append("   - Не находятся выключатели")
report.append("   - Не идентифицируются коробки и датчики")
report.append("")
report.append("3. **Кабельная продукция (0% совпадений)**")
report.append("   - Нет точных совпадений по маркам кабелей")
report.append("   - Возможно, кабели есть, но с другими названиями")
report.append("   - 12 лишних позиций по кабелям — вероятно, неправильно идентифицированы")
report.append("")
report.append("4. **Лотки (100% отсутствует)**")
report.append("   - Файлы 019-025 (Лотки) возвращают 0 элементов")
report.append("   - Легенда не распознается на этих листах")
report.append("   - Отсутствуют все позиции: лотки, углы, консоли, подвесы")
report.append("")
report.append("5. **ПВХ изделия (100% отсутствует)**")
report.append("   - Гофротрубы не находятся")
report.append("   - Металлоконструкции отсутствуют")
report.append("")
report.append("6. **Пусконаладка (100% отсутствует)**")
report.append("   - Все 6 позиций ПНР отсутствуют")
report.append("")

report.append("### 3.3 ⚠️  Проблемы и аномалии")
report.append("")
report.append("1. **Большие расхождения в количествах**")
report.append("   - Некоторые позиции завышены на 700-1066%")
report.append("   - Возможно, дублирование при агрегации")
report.append("   - Или неправильное маппирование вариантов (АЭ/А)")
report.append("")
report.append("2. **Лишние позиции (16 шт)**")
report.append("   - 12 кабельных позиций — возможно, из схем")
report.append("   - 3 светотехнических — дубликаты?")
report.append("   - Требуется анализ причин")
report.append("")
report.append("3. **Наименования в Excel выглядят как номера строк**")
report.append("   - В разделе \"Только в эталоне\" видны цифры 4, 5, 6...")
report.append("   - Возможно, проблема парсинга эталонного DOCX")
report.append("   - Нужна проверка скрипта `vor_compare_excel.py`")
report.append("")

report.append("---")
report.append("")

report.append("## 4. Выводы и рекомендации")
report.append("")
report.append("### 4.1 Текущее состояние")
report.append("")
report.append(f"**Общая точность: 8.1%** (10 совпадений из 123 позиций эталона)")
report.append("")
report.append("Система **частично работает** только для:")
report.append("- Планов освещения (наилучший результат)")
report.append("- Минимально — для схем")
report.append("")
report.append("Система **НЕ работает** для:")
report.append("- Щитового оборудования")
report.append("- Установочных изделий")
report.append("- Лотков и аксессуаров")
report.append("- ПНР и прочих работ")
report.append("")

report.append("### 4.2 Критические проблемы (требуют немедленного исправления)")
report.append("")
report.append("1. **Легенда не распознается на листах лотков (019-025)**")
report.append("   - Диагностика: запустить `diagnose_legend.py` на эти файлы")
report.append("   - Возможно, другой формат таблицы или отсутствие заголовка")
report.append("")
report.append("2. **Щиты не обрабатываются**")
report.append("   - Требуется специальная обработка однолинейных схем")
report.append("   - Возможно, нужен отдельный модуль для `pdf_count_panels.py`")
report.append("")
report.append("3. **Установочные изделия не находятся**")
report.append("   - Проверить, распознается ли легенда на планах привязки")
report.append("   - Возможно, текстовые метки не в том формате")
report.append("")
report.append("4. **Кабельная продукция — нет совпадений**")
report.append("   - Проверить маппинг в `vor_work_mapping.py`")
report.append("   - Возможно, нужно добавить синонимы марок кабелей")
report.append("")

report.append("### 4.3 Рекомендации по доработке")
report.append("")
report.append("**Приоритет 1 (критично):**")
report.append("1. Исправить распознавание легенды на листах лотков")
report.append("2. Добавить обработку однолинейных схем для щитов")
report.append("3. Улучшить извлечение установочных изделий")
report.append("")
report.append("**Приоритет 2 (важно):**")
report.append("4. Проверить и исправить маппинг кабельной продукции")
report.append("5. Исследовать причины завышения количеств (×7-10)")
report.append("6. Добавить генерацию позиций ПНР по найденному оборудованию")
report.append("")
report.append("**Приоритет 3 (желательно):**")
report.append("7. Улучшить скорость обработки (сейчас 18 сек/файл)")
report.append("8. Добавить прогресс-бар с процентами при CLI-запуске")
report.append("9. Исправить парсинг эталонного DOCX в `vor_compare_excel.py`")
report.append("")

report.append("### 4.4 Следующие шаги")
report.append("")
report.append("1. **Детальная диагностика** — запустить на проблемных файлах:")
report.append("   ```bash")
report.append("   python diagnose_legend.py \"Data/.../019-Лотки отм. 0.000.pdf\"")
report.append("   python pdf_overlay.py \"Data/.../005-Планы освещения-отм. 0.000.pdf\"")
report.append("   ```")
report.append("")
report.append("2. **Анализ легенд** — сравнить структуру легенд разных типов листов")
report.append("")
report.append("3. **Тестирование на других захватках** — проверить, воспроизводится ли проблема")
report.append("")
report.append("4. **Улучшение vor_work_mapping** — добавить недостающие маппинги")
report.append("")

report.append("---")
report.append("")
report.append("## 5. Приложения")
report.append("")
report.append("### Файлы результатов")
report.append("")
report.append(f"- Сгенерированный ВОР: `VOR_generated_test.docx`")
report.append(f"- Сравнение в Excel: `VOR_COMPARISON.xlsx`")
report.append(f"- Этот отчет: `TEST_REPORT.md`")
report.append("")
report.append("### Команды для воспроизведения")
report.append("")
report.append("```bash")
report.append("# Генерация ВОР")
report.append("python test_gpk3_vor.py")
report.append("")
report.append("# Сравнение с эталоном")
report.append("python vor_compare_excel.py \"ВОР ЭО, Захватка 3_ГПК.docx\" \"VOR_generated_test.docx\" -o VOR_COMPARISON.xlsx")
report.append("")
report.append("# Анализ результатов")
report.append("python generate_test_report.py")
report.append("```")
report.append("")

# Save report
REPORT_FILE.write_text("\n".join(report), encoding="utf-8")

print(f"✅ Отчет сохранен: {REPORT_FILE}")
print()
print("\n".join(report))
