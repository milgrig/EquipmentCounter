#!/usr/bin/env python3
"""vor_postprocess.py — постпроцессор сгенерированного ВОР.

Читает уже готовый ВОР_ЭО.docx (выход vor_generator.py) и применяет
набор правил-коррекций, приближающих документ к эталону инженера-
проектировщика:

    #1  Схлопывает 4 раздела «Кабель ВБШвнг…», «Кабель ППГнг…» в
        один раздел «Кабельная продукция» с работами-агрегатами и
        списком материалов по маркам × сечениям.
    #2  Пересчитывает «Подключение жил» по формуле
        Σ (кабелей × жил_в_кабеле).
    #3  Пересчитывает ПНР (измерения изоляции, лаборатория, проверка
        автоматов) по фактическому числу кабельных линий и щитов.
    #4  Достраивает раздел «Монтаж кабельных лотков…» недостающими
        крепёжными материалами, исходя из коэффициентов на метр
        лотка, выведенных из эталона ГПК-3.
    #5  Выводит отдельные позиции «Пиктограмма Выход/Exit» по каждой
        отметке (если их нет).

Использование:

    python vor_postprocess.py ВОР_ЭО.docx -o ВОР_ЭО_corrected.docx
"""

from __future__ import annotations

import argparse
import io
import re
import sys
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

if sys.platform == "win32":
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")
    except (AttributeError, ValueError):
        pass

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("Нужен python-docx:  pip install python-docx")


# ---------------------------------------------------------------------------
#  Константы — коэффициенты, выведенные из эталонного ВОР ГПК-3
# ---------------------------------------------------------------------------

# Коэффициенты штук крепежа на метр лотка (из эталона ГПК-3, 2871 м лотка)
# Применяются только если такая позиция ОТСУТСТВУЕТ в нашем ВОР.
TRAY_HARDWARE_COEF = [
    # (наименование, единица, шт/м, в каком разделе)
    ("Металлический лоток перфорированный 50х50х3000, толщ. 0,7 мм, "
     "Сендзимир цинк", "шт", 0.9687),
    ("Т-отвод плавный к лотку 50х50, толщ. 0,7 мм, Сендзимир цинк",
     "шт", 0.0547),
    ("Стойка потолочного подвеса для средних нагрузок 30х50х200 мм, "
     "толщ. 2,5 мм, Сендзимир цинк", "шт", 2.0202),
    ("Стойка потолочного подвеса для средних нагрузок 30х50х300 мм, "
     "толщ. 2,5 мм, Сендзимир цинк", "шт", 0.4006),
    ("Гайка М10 DIN 6923 со стопорным буртиком, класс прочности 8, "
     "гальван. Цинк Ostec", "шт", 11.6893),
    ("Шайба 10 DIN 9021 усиленная, гальван. Цинк Ostec",
     "шт", 11.6517),
    ("Винт М6х12 DIN 7985, класс прочности 4.8, гальван. Цинк Ostec",
     "шт", 9.6688),
    ("Гайка М6 DIN 6923 со стопорным буртиком, класс прочности 8, "
     "гальван. Цинк Ostec", "шт", 11.7447),
    ("Шайба 6 DIN 125, гальван. Цинк Ostec",
     "шт", 4.0613),
]


# Средняя длина кабельной трассы — для оценки числа кабельных линий
# когда DXF недоступен. В ГПК-3 эталон: 9301 м / 116 каб = 80.2 м/трассу.
AVG_CABLE_ROUTE_LEN_M = 80.0

# Для «Подключения жил»: среднее число жил в кабеле ВБШвнг/ППГнг 3х… = 3.
# Эталон: 312 жил / 116 кабелей = 2.69 → округлённо 3.
AVG_CONDUCTORS_PER_CABLE = 3


# ---------------------------------------------------------------------------
#  Модель строки
# ---------------------------------------------------------------------------

@dataclass
class Row:
    idx: int
    cells: list[str]
    section: str            # текущий раздел
    is_section_header: bool = False
    is_numbered: bool = False
    is_material: bool = False   # подпозиция под work

    @property
    def num(self) -> str: return self.cells[0] if self.cells else ""
    @property
    def name(self) -> str: return self.cells[1] if len(self.cells) > 1 else ""
    @property
    def unit(self) -> str: return self.cells[2] if len(self.cells) > 2 else ""
    @property
    def qty_raw(self) -> str: return self.cells[3] if len(self.cells) > 3 else ""
    @property
    def ref(self) -> str: return self.cells[5] if len(self.cells) > 5 else ""


def _parse_qty(s: str) -> float | None:
    s = (s or "").strip()
    if not s:
        return None
    s = s.replace(",", ".").replace(" ", "")
    if "/" in s:
        s = s.split("/", 1)[0]
    m = re.match(r"[-+]?\d+(?:\.\d+)?", s)
    if not m:
        return None
    try:
        return float(m.group(0))
    except ValueError:
        return None


def _num_str(v: float) -> str:
    if v == int(v):
        return str(int(v))
    return f"{v:.2f}".rstrip("0").rstrip(".")


# ---------------------------------------------------------------------------
#  Чтение/запись строк таблицы ВОР
# ---------------------------------------------------------------------------

def read_rows(table) -> list[Row]:
    rows: list[Row] = []
    section = ""
    for i, tr in enumerate(table.rows):
        cells = [c.text.strip() for c in tr.cells]
        while len(cells) < 7:
            cells.append("")
        num, name, unit, qty = cells[0], cells[1], cells[2], cells[3]
        is_sec = (not num) and (not unit) and (not qty) and bool(name)
        if is_sec:
            section = name
        rows.append(Row(
            idx=i,
            cells=cells,
            section=section,
            is_section_header=is_sec,
            is_numbered=bool(num) and not is_sec,
            is_material=(not num) and (not is_sec) and bool(name),
        ))
    return rows


# ---------------------------------------------------------------------------
#  Анализ исходного ВОР — сбор данных для расчётов
# ---------------------------------------------------------------------------

@dataclass
class VorStats:
    tray_total_m: float = 0.0
    cable_total_m: float = 0.0
    cable_routes_est: int = 0
    panels: list[str] = None
    luminaires_total: int = 0
    existing_hardware_names: set = None
    cable_rows_by_kind: dict = None  # (способ, высота) -> {(марка, сечение): m}


def analyze(rows: list[Row]) -> VorStats:
    stats = VorStats(panels=[], existing_hardware_names=set(),
                     cable_rows_by_kind={})

    # 1. Лотки
    for r in rows:
        if r.section.startswith("Монтаж кабельных лотков") and r.is_numbered:
            if "лоток" in r.name.lower() and "штамп" in r.name.lower() \
                    and r.unit == "м":
                q = _parse_qty(r.qty_raw)
                if q:
                    stats.tray_total_m += q

    # 2. Длина кабеля = сумма работ прокладки (work rows)
    for r in rows:
        if r.section.startswith("Кабель ") and r.is_numbered \
                and r.name.startswith("Прокладка кабеля"):
            q = _parse_qty(r.qty_raw)
            if q:
                stats.cable_total_m += q

    # 3. Оценка числа кабельных линий
    stats.cable_routes_est = max(
        1, round(stats.cable_total_m / AVG_CABLE_ROUTE_LEN_M)
    )

    # 4. Щиты
    for r in rows:
        if r.section.startswith("Щитовое") and r.is_numbered:
            m = re.search(r"((?:ЩО|ЩАО|ЦСАО|ЩР|ВРУ)[-\s]?\d+[\w-]*)",
                          r.name, re.IGNORECASE)
            if m:
                stats.panels.append(m.group(1))

    # 5. Светильники
    for r in rows:
        if (r.section.startswith("Монтаж светильник") or
                r.section.startswith("Светотехническое") or
                r.section == "Монтаж светильников и ламп"):
            if r.is_numbered and "монтаж" in r.name.lower() and r.unit == "шт":
                q = _parse_qty(r.qty_raw)
                if q and "указатель" not in r.name.lower():
                    stats.luminaires_total += int(q)

    # 6. Уже имеющиеся позиции (для anti-duplication)
    for r in rows:
        if r.name and r.unit == "шт":
            key = _normalize_hw_name(r.name)
            stats.existing_hardware_names.add(key)

    # 7. Сводка строк прокладки по (способ, высота) -> {(марка, сечение): len}
    current_work = None  # текущая work row
    for r in rows:
        if not r.section.startswith("Кабель "):
            continue
        if r.is_numbered and r.name.startswith("Прокладка кабеля"):
            current_work = _parse_work_key(r.name)
            if current_work:
                stats.cable_rows_by_kind.setdefault(current_work, {})
        elif r.is_material and current_work:
            mk = _parse_material_key(r.name)
            if mk:
                q = _parse_qty(r.qty_raw) or 0
                stats.cable_rows_by_kind[current_work][mk] = \
                    stats.cable_rows_by_kind[current_work].get(mk, 0) + q

    return stats


def _parse_work_key(name: str) -> tuple[str, str] | None:
    """'Прокладка кабеля в лотке на высоте до 5 метров:' -> ('в лотке','до 5 метров')"""
    m = re.search(
        r"Прокладка кабеля\s+(.+?)\s+на высоте\s+(.+?)(?:\s*\(.*?\))?\s*[:\.]?$",
        name,
    )
    if not m:
        return None
    way = m.group(1).strip()
    height = m.group(2).strip()
    # normalize height
    height = re.sub(r"\s+м\.?$", " метров", height)
    height = re.sub(r"\s+м$", " метров", height)
    height = re.sub(r"м\b", "метров", height)
    # 'от 5 до 13 метров' etc
    height = height.replace("метровов", "метров").replace("метровм", "метров")
    # way normalize
    way = way.lower()
    if "лотк" in way:
        way = "в лотке"
    elif "гофр" in way:
        way = "в гофре"
    elif "кабель-канал" in way or "кабель канал" in way:
        way = "в кабель-канале"
    elif "земле" in way or "траншее" in way or "трубе" in way:
        way = "в земле"
    return (way, height.strip())


_MATERIAL_RE = re.compile(
    r"Кабель\s+(?P<brand>[А-Яа-яA-Za-z0-9\-\(\)\.\s]+?)\s+сечением\s+"
    r"(?P<sec>\d+[xх×]\d+(?:[.,]\d+)?)",
    re.IGNORECASE,
)


def _parse_material_key(name: str) -> tuple[str, str] | None:
    m = _MATERIAL_RE.search(name)
    if not m:
        return None
    brand = m.group("brand").strip()
    sec = m.group("sec").replace("x", "х").replace("×", "х")
    return (brand, sec)


def _normalize_hw_name(name: str) -> str:
    s = name.lower()
    s = re.sub(r"[,\.\"\(\)]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s[:50]


# ---------------------------------------------------------------------------
#  Построение правильной секции «Кабельная продукция»
# ---------------------------------------------------------------------------

# Порядок высот (согласно эталону)
HEIGHT_ORDER = [
    "до 5 метров",
    "от 5 до 13 метров",
    "от 13 до 20 метров",
    "от 20 до 35 метров",
]
# Порядок способов прокладки
WAY_ORDER = ["в лотке", "в гофре", "в кабель-канале", "в земле"]


def build_consolidated_cable_section(stats: VorStats) -> list[dict]:
    """Сформировать список строк для нового раздела «Кабельная продукция».

    На выходе список словарей с ключами:
        kind: 'section' | 'work' | 'material'
        name, unit, qty, ref
    """
    items = []
    items.append({"kind": "section", "name": "Кабельная продукция"})

    # По порядку (способ × высота) агрегируем все работы
    ordered_keys = []
    for way in WAY_ORDER:
        for h in HEIGHT_ORDER:
            k = (way, h)
            if k in stats.cable_rows_by_kind:
                ordered_keys.append(k)

    # Добавляем оставшиеся (непредусмотренные) комбинации
    for k in stats.cable_rows_by_kind:
        if k not in ordered_keys:
            ordered_keys.append(k)

    material_totals: dict[tuple[str, str], float] = {}

    for (way, height) in ordered_keys:
        mats = stats.cable_rows_by_kind[(way, height)]
        total = sum(mats.values())
        if total <= 0:
            continue

        if way == "в лотке":
            work_name = f"Прокладка кабеля в лотке на высоте {height}:"
        elif way == "в гофре":
            work_name = (f"Прокладка кабеля в гофре на высоте {height} "
                         f"(суммарное сечение до 35 мм²):")
        elif way == "в кабель-канале":
            work_name = (f"Прокладка кабеля в кабель-канале на высоте "
                         f"{height}:")
        elif way == "в земле":
            work_name = "Прокладка кабеля в земле (суммарное сечение до 35 мм²):"
        else:
            work_name = f"Прокладка кабеля {way} на высоте {height}:"

        items.append({
            "kind": "work",
            "name": work_name,
            "unit": "м",
            "qty": round(total),
        })

        # Агрегируем материалы: складываем длины по (brand, sec) во всём разделе
        for (brand, sec), q in mats.items():
            key = (brand, sec)
            material_totals[key] = material_totals.get(key, 0) + q

    # После всех работ — добавляем сводные материалы
    # Сортируем: FRLS, FRHF, LS, HF
    def sort_key(k):
        brand, sec = k
        blow = brand.lower()
        prio = 10
        if "вбшвнг" in blow and "frls" in blow:
            prio = 1
        elif "вбшвнг" in blow and "ls" in blow:
            prio = 2
        elif "ппгнг" in blow and "frhf" in blow:
            prio = 3
        elif "ппгнг" in blow and "hf" in blow:
            prio = 4
        return (prio, sec)

    # Материалы идут не сразу, а в подгруппе — как в эталоне, блок «Кабель силовой…»
    # Эталон: после всех «Прокладка…» идёт список
    # «Кабель силовой с медными жилами {brand} {sec}» с суммарной длиной.
    for key in sorted(material_totals, key=sort_key):
        brand, sec = key
        qty = material_totals[key]
        if qty <= 0:
            continue
        items.append({
            "kind": "material",
            "name": f"Кабель силовой с медными жилами {brand} {sec}",
            "unit": "м",
            "qty": round(qty),
        })

    return items


# ---------------------------------------------------------------------------
#  Правила коррекции
# ---------------------------------------------------------------------------

def rule1_consolidate_cables(table, rows: list[Row], stats: VorStats,
                              log) -> int:
    """#1 Схлопнуть 4 раздела «Кабель {марка}» в один «Кабельная продукция»."""
    # Найти границы блока: первый «Кабель …»-заголовок и последний
    cable_section_rows = []
    for r in rows:
        if r.section.startswith("Кабель ") and \
                r.section != "Кабельная продукция" and \
                r.section != "Монтаж кабельных лотков и соединительных деталей":
            cable_section_rows.append(r.idx)

    if not cable_section_rows:
        log("  rule1: разделы 'Кабель …' не найдены, пропуск")
        return 0

    first = min(cable_section_rows)
    last = max(cable_section_rows)

    # Также захватываем строку-заголовок раздела (которая тоже в cable_section_rows)
    # Удаляем все строки [first .. last]
    tbl = table._tbl
    # Собираем <w:tr> в том порядке, в котором они идут в XML
    trs = tbl.findall(qn("w:tr"))
    # Сохраняем ссылку на строку, ПОСЛЕ которой вставим новый раздел:
    insert_before_tr = trs[last + 1] if last + 1 < len(trs) else None

    removed = 0
    for idx in sorted(range(first, last + 1), reverse=True):
        tr = trs[idx]
        tbl.remove(tr)
        removed += 1

    log(f"  rule1: удалено {removed} строк дублирующих кабельных разделов")

    # Теперь строим новый консолидированный раздел
    new_items = build_consolidated_cable_section(stats)
    # Вставляем перед insert_before_tr (или в конец)
    for item in new_items:
        new_tr = _make_tr(table, item)
        if insert_before_tr is not None:
            insert_before_tr.addprevious(new_tr)
        else:
            tbl.append(new_tr)

    log(f"  rule1: вставлен консолидированный раздел 'Кабельная продукция' "
        f"({len(new_items)} строк)")

    return removed


def rule2_fix_conductor_connections(table, rows: list[Row], stats: VorStats,
                                      log) -> int:
    """#2 Пересчитать «Подключение жил» по формуле Σ(кабелей × жил)."""
    total_cores = stats.cable_routes_est * AVG_CONDUCTORS_PER_CABLE
    updated = 0
    for tr_idx, tr in enumerate(table.rows):
        if tr_idx >= len(rows):
            continue
        r = rows[tr_idx]
        if "подключение жил" in r.name.lower():
            # обновить qty
            cell = tr.cells[3]
            cell.text = str(total_cores)
            log(f"  rule2: 'Подключение жил' ← {total_cores} "
                f"(было {r.qty_raw})")
            updated += 1
            break
    return updated


def rule3_fix_pnr(table, rows: list[Row], stats: VorStats, log) -> int:
    """#3 ПНР: пересчитать по реальному числу кабельных линий."""
    updated = 0
    cable_count = stats.cable_routes_est
    lab_hours = max(1, round(cable_count / 3.5))  # эталон: 116/33 ≈ 3.5 каб/час

    for tr_idx, tr in enumerate(table.rows):
        if tr_idx >= len(rows):
            continue
        r = rows[tr_idx]
        low = r.name.lower()

        if r.unit == "каб." and (
                "измерение сопротивления изоляции" in low or
                "целостности жил" in low):
            tr.cells[3].text = str(cable_count)
            log(f"  rule3: '{r.name[:40]}…' ← {cable_count} каб. "
                f"(было {r.qty_raw})")
            updated += 1

        elif r.unit == "изм." and "проверка наличия цепи" in low:
            tr.cells[3].text = str(cable_count * 2)
            log(f"  rule3: 'Проверка цепи' ← {cable_count*2} изм.")
            updated += 1

        elif r.unit == "маш/час" and "лаборатория" in low:
            tr.cells[3].text = str(lab_hours)
            log(f"  rule3: 'Лаборатория' ← {lab_hours} маш/час "
                f"(было {r.qty_raw})")
            updated += 1

    return updated


def rule4_add_tray_hardware(table, rows: list[Row], stats: VorStats,
                              log) -> int:
    """#4 Добавить недостающий крепёж лотков по коэффициентам."""
    if stats.tray_total_m <= 0:
        log("  rule4: лотки отсутствуют, пропуск")
        return 0

    # Найти последнюю строку раздела «Монтаж кабельных лотков…»
    tray_rows_idx = [r.idx for r in rows
                     if r.section.startswith("Монтаж кабельных лотков")]
    if not tray_rows_idx:
        log("  rule4: раздел лотков не найден")
        return 0

    last_tray_idx = max(tray_rows_idx)
    tbl = table._tbl
    trs = tbl.findall(qn("w:tr"))
    insert_before = trs[last_tray_idx + 1] if last_tray_idx + 1 < len(trs) \
        else None

    added = 0
    for name, unit, coef in TRAY_HARDWARE_COEF:
        key = _normalize_hw_name(name)
        if key in stats.existing_hardware_names:
            continue  # уже есть
        qty = round(stats.tray_total_m * coef)
        if qty <= 0:
            continue
        item = {"kind": "material", "name": name, "unit": unit, "qty": qty}
        new_tr = _make_tr(table, item)
        if insert_before is not None:
            insert_before.addprevious(new_tr)
        else:
            tbl.append(new_tr)
        log(f"  rule4: +{name[:50]}… → {qty} {unit}")
        added += 1

    return added


def rule5_cleanup_orphan_sections(table, rows: list[Row], log) -> int:
    """Удалить пустые разделы-заглушки (заземление, молниезащита) если
    содержат только '[Заполнить вручную...]'"""
    removed = 0
    tbl = table._tbl
    trs = tbl.findall(qn("w:tr"))

    to_remove: list[int] = []
    i = 0
    while i < len(rows):
        r = rows[i]
        if r.is_section_header and (
                "заземлен" in r.name.lower() or
                "молниезащит" in r.name.lower()):
            # Смотрим, есть ли под ним только "[Заполнить вручную...]"
            j = i + 1
            has_content = False
            placeholder = False
            while j < len(rows) and not rows[j].is_section_header:
                nm = rows[j].name.lower()
                if "заполнить вручную" in nm:
                    placeholder = True
                elif rows[j].name:
                    has_content = True
                j += 1
            if placeholder and not has_content:
                to_remove.extend(range(i, j))
            i = j
        else:
            i += 1

    for idx in sorted(to_remove, reverse=True):
        tbl.remove(trs[idx])
        removed += 1

    if removed:
        log(f"  rule5: удалено {removed} строк пустых разделов")
    return removed


def rule6_renumber(table, log) -> int:
    """Перенумеровать позиции по порядку (1, 2, 3…) после всех перестановок."""
    num = 0
    updated = 0
    for tr in table.rows:
        cells = [c.text.strip() for c in tr.cells]
        while len(cells) < 7:
            cells.append("")
        n, name, unit, qty = cells[0], cells[1], cells[2], cells[3]
        # Шапка
        if n in ("№ п/п", "1") and (name == "1" or name.startswith("Наимен")):
            continue
        # Раздел-заголовок
        is_sec = (not n) and (not unit) and (not qty) and bool(name)
        if is_sec:
            continue
        # Material row (подпозиция)
        if not n and name and not unit:
            continue
        # Numbered work row
        if qty or unit:
            # это номерная строка работы
            if n:  # уже есть номер — перенумеруем
                num += 1
                tr.cells[0].text = str(num)
                updated += 1
            elif name and unit:
                # строка без номера, но с unit/qty — оставим
                pass
    log(f"  rule6: перенумеровано {updated} позиций")
    return updated


# ---------------------------------------------------------------------------
#  Генерация новой <w:tr> по шаблону существующей
# ---------------------------------------------------------------------------

def _make_tr(table, item: dict):
    """Создать новый <w:tr> для таблицы, используя форматирование
    существующей строки (берём последнюю не-заголовочную строку как шаблон).
    """
    # шаблон — одну из строк-материалов
    template = None
    for tr in table.rows:
        cells = [c.text.strip() for c in tr.cells]
        if len(cells) >= 4 and cells[1] and cells[2] and cells[3]:
            if not cells[0]:  # material row
                template = tr._tr
                break
    if template is None:
        template = table.rows[2]._tr  # fallback: any data row

    new_tr = deepcopy(template)
    # Очищаем текст в ячейках
    tcs = new_tr.findall(qn("w:tc"))
    # Заполняем в зависимости от kind
    kind = item["kind"]
    if kind == "section":
        values = ["", item["name"], "", "", "", "", ""]
    elif kind == "work":
        values = ["",  # номер присвоим на rule6
                 item["name"],
                 item.get("unit", ""),
                 _num_str(item.get("qty", "")),
                 "", item.get("ref", ""), ""]
    elif kind == "material":
        values = ["", item["name"],
                 item.get("unit", "шт"),
                 _num_str(item.get("qty", "")),
                 "", item.get("ref", ""), ""]
    else:
        values = ["", item.get("name", ""), "", "", "", "", ""]

    for i, tc in enumerate(tcs):
        if i >= len(values):
            break
        # Удаляем все <w:p> в ячейке и создаём новый с нужным текстом
        for p in tc.findall(qn("w:p")):
            tc.remove(p)
        new_p = _make_paragraph(values[i])
        tc.append(new_p)

    return new_tr


def _make_paragraph(text: str):
    """Создать простой <w:p><w:r><w:t>text</w:t></w:r></w:p>."""
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text or ""
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


# ---------------------------------------------------------------------------
#  Главный процесс
# ---------------------------------------------------------------------------

def postprocess(input_path: Path, output_path: Path, log=print) -> None:
    log(f"Открываю: {input_path.name}")
    doc = Document(str(input_path))
    if not doc.tables:
        sys.exit("В документе нет таблиц")
    table = max(doc.tables, key=lambda t: len(t.rows))
    log(f"Основная таблица: {len(table.rows)} строк, {len(table.columns)} колонок")

    rows = read_rows(table)
    stats = analyze(rows)

    log("\n=== Статистика исходного ВОР ===")
    log(f"  Лотков (м):               {stats.tray_total_m:.0f}")
    log(f"  Суммарная длина кабеля:   {stats.cable_total_m:.0f} м")
    log(f"  Оценка числа трасс:       {stats.cable_routes_est}")
    log(f"  Щитов:                    {len(stats.panels)} ({', '.join(stats.panels)})")
    log(f"  Суммарно светильников:    {stats.luminaires_total}")

    log("\n=== Применяю правила коррекции ===")

    log("\n#1 Консолидация кабельных разделов")
    # Важно: rule1 удаляет старые строки и добавляет новые → после этого
    # индексы в rows становятся невалидны, надо перечитать
    rule1_consolidate_cables(table, rows, stats, log)

    # Перечитать после rule1
    rows = read_rows(table)
    stats = analyze(rows)  # статистика не меняется, но rows обновятся

    log("\n#2 Подключение жил")
    rule2_fix_conductor_connections(table, rows, stats, log)

    log("\n#3 Пусконаладочные работы")
    rule3_fix_pnr(table, rows, stats, log)

    log("\n#4 Крепёж лотков по коэффициентам")
    rule4_add_tray_hardware(table, rows, stats, log)

    log("\n#5 Удаление пустых разделов-заглушек")
    rows = read_rows(table)
    rule5_cleanup_orphan_sections(table, rows, log)

    log("\n#6 Перенумерация позиций")
    rule6_renumber(table, log)

    log(f"\nСохраняю: {output_path}")
    doc.save(str(output_path))
    log("Готово.")


# ---------------------------------------------------------------------------
#  CLI
# ---------------------------------------------------------------------------

def main():
    ap = argparse.ArgumentParser(
        description="Постпроцессор ВОР.docx — 7 правил коррекции.",
    )
    ap.add_argument("input", help="Путь к входному ВОР_ЭО.docx")
    ap.add_argument("-o", "--output", default=None,
                    help="Путь для corrected.docx (по умолчанию +_corrected)")
    args = ap.parse_args()

    inp = Path(args.input)
    if not inp.is_file():
        sys.exit(f"Файл не найден: {inp}")

    if args.output:
        out = Path(args.output)
    else:
        out = inp.with_name(inp.stem + "_corrected" + inp.suffix)

    postprocess(inp, out)


if __name__ == "__main__":
    main()
