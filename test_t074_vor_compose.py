#!/usr/bin/env python3
"""test_t074_vor_compose.py -- Verify composed VOR against reference (T074 / S016).

Task: Run full pipeline on 3-zahvatka folder using new vor_compose,
compare to reference "VOR EO Zahvatka 3 GPK.docx" using diff algorithm
similar to _compare_two_vors.py.

Acceptance Gates:
  G1: row-match rate >= 65% (was 25.4%)
  G2: cable total qty within 50% of 14000m reference (target 7000-21000m)
  G3: luminaire total qty within 30% of reference ~1100 sht (target 770-1430)
  G4: every reference row category (luminaire, cable, tray, switch, post,
      junction-box) has at least one row in generated output

Output: discussions/T067_verification.md
"""

from __future__ import annotations

import io
import logging
import re
import sys
from collections import defaultdict
from pathlib import Path
from typing import Any

# Project root
PROJECT_ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(PROJECT_ROOT))

# Import pipeline components
from web_app import _count_equipment_in_pdf
from vor_compose import compose_vor_table

# Import docx reader
from docx import Document

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger("test_t074")

# Configure output encoding
sys.stdout.reconfigure(encoding='utf-8')

# ============================================================================
# Configuration
# ============================================================================

PDF_DIR = Path(r"Data\ДБТ разделы для ИИ\03_ГПК_\3-я захватка\02_PDF")
REFERENCE_DOCX = Path(r"Data\ДБТ разделы для ИИ\03_ГПК_\3-я захватка\ВОР ЭО, Захватка 3_ГПК.docx")
OUTPUT_REPORT = Path(r".tayfa\common\discussions\T067_verification.md")

# Acceptance gate thresholds
GATE_G1_ROW_MATCH_MIN = 65.0  # percent
GATE_G2_CABLE_MIN = 7000.0    # meters
GATE_G2_CABLE_MAX = 21000.0   # meters
GATE_G3_LUMINAIRE_MIN = 770   # count
GATE_G3_LUMINAIRE_MAX = 1430  # count

# Categories for G4 check
REQUIRED_CATEGORIES = [
    "luminaire",
    "cable",
    "tray",
    "switch",
    "post",
    "junction-box"
]


# ============================================================================
# Reference VOR parsing
# ============================================================================

def parse_reference_vor(docx_path: Path) -> list[dict]:
    """Parse reference DOCX VOR table into list of row dicts.

    Returns:
        List of dicts with keys: name, unit, qty, formula, drawing_refs, notes
    """
    doc = Document(str(docx_path))

    # Find the main VOR table (largest table)
    if not doc.tables:
        log.error("No tables found in reference DOCX")
        return []

    main_table = max(doc.tables, key=lambda t: len(t.rows))
    log.info("Reference table has %d rows, %d columns",
             len(main_table.rows), len(main_table.columns))

    rows = []
    header_row_count = 0

    for ri, row in enumerate(main_table.rows):
        cells = [c.text.strip() for c in row.cells]

        # Skip header rows (look for column labels)
        if ri < 3:
            text = " ".join(cells).lower()
            if "наименование" in text or "ед. изм" in text or "п/п" in text:
                header_row_count = ri + 1
                continue

        if ri < header_row_count:
            continue

        # Skip empty rows
        if not any(cells):
            continue

        # Parse data row (assuming 7 columns: #, Name, Unit, Qty, Formula, Refs, Notes)
        if len(cells) >= 4:
            try:
                # Column 1: row number (optional)
                # Column 2: name
                # Column 3: unit
                # Column 4: qty
                name = cells[1] if len(cells) > 1 else ""
                unit = cells[2] if len(cells) > 2 else ""
                qty_str = cells[3] if len(cells) > 3 else "0"
                formula = cells[4] if len(cells) > 4 else ""
                refs = cells[5] if len(cells) > 5 else ""
                notes = cells[6] if len(cells) > 6 else ""

                # Parse quantity
                qty = 0.0
                try:
                    # Remove spaces, replace comma with dot
                    qty_clean = qty_str.replace(" ", "").replace(",", ".")
                    qty = float(qty_clean)
                except (ValueError, AttributeError):
                    qty = 0.0

                if name and qty > 0:
                    rows.append({
                        "name": name,
                        "unit": unit,
                        "qty": qty,
                        "formula": formula,
                        "drawing_refs": refs,
                        "notes": notes
                    })
            except Exception as e:
                log.warning("Failed to parse row %d: %s", ri, e)
                continue

    log.info("Parsed %d data rows from reference VOR", len(rows))
    return rows


# ============================================================================
# Generated VOR processing
# ============================================================================

def generate_vor_from_pdfs(pdf_dir: Path) -> list[dict]:
    """Generate VOR table using vor_compose pipeline.

    Returns:
        List of composed VOR row dicts
    """
    log.info("="*80)
    log.info("Generating VOR from PDFs in: %s", pdf_dir)
    log.info("="*80)

    # Find all PDF files (excluding the reference VOR PDF itself)
    pdf_files = sorted([
        p for p in pdf_dir.glob("*.pdf")
        if "ВОР" not in p.name  # Skip reference VOR PDF
    ])

    log.info("Found %d PDF files", len(pdf_files))

    # Count equipment in each PDF
    items_per_pdf: dict[str, list[dict]] = {}

    for pdf_path in pdf_files:
        log.info("Processing: %s", pdf_path.name)
        try:
            file_result = _count_equipment_in_pdf(str(pdf_path))
            items_per_pdf[pdf_path.name] = file_result
            log.info("  -> %d items", len(file_result))
        except Exception as exc:
            log.error("  -> Failed: %s", exc, exc_info=True)
            items_per_pdf[pdf_path.name] = []

    # Compose VOR table
    log.info("\nComposing VOR table...")
    vor_rows = compose_vor_table(items_per_pdf)
    log.info("Composed %d VOR rows", len(vor_rows))

    return vor_rows


# ============================================================================
# Comparison and analysis
# ============================================================================

def normalize_name(name: str) -> str:
    """Normalize a VOR row name for fuzzy matching."""
    # Convert to lowercase
    text = name.lower()
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    # Remove punctuation
    text = re.sub(r'[^\w\s\u0400-\u04FF]', '', text)
    return text


def calculate_name_similarity(name1: str, name2: str) -> float:
    """Calculate similarity score between two names (0.0 to 1.0).

    Uses simple token overlap method.
    """
    tokens1 = set(normalize_name(name1).split())
    tokens2 = set(normalize_name(name2).split())

    if not tokens1 or not tokens2:
        return 0.0

    intersection = tokens1 & tokens2
    union = tokens1 | tokens2

    return len(intersection) / len(union)


def match_vor_rows(reference: list[dict], generated: list[dict],
                   threshold: float = 0.6) -> tuple[list[tuple], dict]:
    """Match reference rows to generated rows.

    Returns:
        (matches, stats) where matches is list of (ref_idx, gen_idx, score)
        and stats is a dict with match statistics
    """
    matches = []
    ref_matched = set()
    gen_matched = set()

    # For each reference row, find best matching generated row
    for ri, ref_row in enumerate(reference):
        best_score = 0.0
        best_gi = -1

        for gi, gen_row in enumerate(generated):
            if gi in gen_matched:
                continue

            # Calculate name similarity
            name_sim = calculate_name_similarity(ref_row["name"], gen_row["name"])

            # Bonus if units match
            unit_bonus = 0.1 if ref_row["unit"] == gen_row["unit"] else 0.0

            score = name_sim + unit_bonus

            if score > best_score and score >= threshold:
                best_score = score
                best_gi = gi

        if best_gi >= 0:
            matches.append((ri, best_gi, best_score))
            ref_matched.add(ri)
            gen_matched.add(best_gi)

    stats = {
        "total_reference": len(reference),
        "total_generated": len(generated),
        "matched": len(matches),
        "ref_unmatched": len(reference) - len(ref_matched),
        "gen_unmatched": len(generated) - len(gen_matched),
        "match_rate": 100.0 * len(matches) / len(reference) if reference else 0.0
    }

    return matches, stats


def categorize_row(row: dict) -> str:
    """Categorize a VOR row by detecting keywords in name."""
    name = row["name"].lower()

    # Order matters - check most specific first
    if any(k in name for k in ["светильник", "светодиодн", "luminaire"]):
        return "luminaire"
    elif any(k in name for k in ["кабел", "cable", "провод"]):
        return "cable"
    elif any(k in name for k in ["лоток", "tray"]):
        return "tray"
    elif any(k in name for k in ["выключател", "switch"]):
        return "switch"
    elif any(k in name for k in ["пост управления", "пост управ", "control post"]):
        return "post"
    elif any(k in name for k in ["коробк", "junction", "box"]):
        return "junction-box"
    else:
        return "other"


def calculate_category_totals(rows: list[dict]) -> dict[str, dict]:
    """Calculate quantity totals by category.

    Returns:
        Dict mapping category -> {unit: total_qty}
    """
    totals = defaultdict(lambda: defaultdict(float))

    for row in rows:
        category = categorize_row(row)
        unit = row.get("unit", "")
        qty = row.get("qty", 0.0) or row.get("total", 0.0) or 0.0

        totals[category][unit] += qty

    return totals


def check_gates(reference: list[dict], generated: list[dict],
                matches: list[tuple], stats: dict) -> dict[str, Any]:
    """Check all acceptance gates.

    Returns:
        Dict with gate results and pass/fail status
    """
    results = {}

    # G1: Row-match rate >= 65%
    match_rate = stats["match_rate"]
    results["G1"] = {
        "description": "Row-match rate >= 65%",
        "value": match_rate,
        "threshold": GATE_G1_ROW_MATCH_MIN,
        "pass": match_rate >= GATE_G1_ROW_MATCH_MIN
    }

    # Calculate category totals
    ref_totals = calculate_category_totals(reference)
    gen_totals = calculate_category_totals(generated)

    # G2: Cable total within 7000-21000m
    ref_cable_m = ref_totals["cable"].get("м", 0.0) + ref_totals["cable"].get("m", 0.0)
    gen_cable_m = gen_totals["cable"].get("м", 0.0) + gen_totals["cable"].get("m", 0.0)

    results["G2"] = {
        "description": "Cable total within 7000-21000m",
        "ref_value": ref_cable_m,
        "gen_value": gen_cable_m,
        "min_threshold": GATE_G2_CABLE_MIN,
        "max_threshold": GATE_G2_CABLE_MAX,
        "pass": GATE_G2_CABLE_MIN <= gen_cable_m <= GATE_G2_CABLE_MAX
    }

    # G3: Luminaire total within 770-1430 sht
    ref_lum_sht = ref_totals["luminaire"].get("шт", 0.0) + ref_totals["luminaire"].get("sht", 0.0)
    gen_lum_sht = gen_totals["luminaire"].get("шт", 0.0) + gen_totals["luminaire"].get("sht", 0.0)

    results["G3"] = {
        "description": "Luminaire total within 770-1430 sht",
        "ref_value": ref_lum_sht,
        "gen_value": gen_lum_sht,
        "min_threshold": GATE_G3_LUMINAIRE_MIN,
        "max_threshold": GATE_G3_LUMINAIRE_MAX,
        "pass": GATE_G3_LUMINAIRE_MIN <= gen_lum_sht <= GATE_G3_LUMINAIRE_MAX
    }

    # G4: All required categories present
    categories_present = {}
    for cat in REQUIRED_CATEGORIES:
        present = any(categorize_row(row) == cat for row in generated)
        categories_present[cat] = present

    all_present = all(categories_present.values())

    results["G4"] = {
        "description": "All required categories present",
        "categories": categories_present,
        "pass": all_present
    }

    return results


# ============================================================================
# Report generation
# ============================================================================

def generate_markdown_report(reference: list[dict], generated: list[dict],
                            matches: list[tuple], stats: dict,
                            gate_results: dict) -> str:
    """Generate markdown report for T067_verification.md."""

    lines = []
    lines.append("# T067/T074 Verification: VOR Composition with vor_compose.py")
    lines.append("")
    lines.append("**Date:** 2026-05-12")
    lines.append("**Sprint:** S016")
    lines.append("**Branch:** sprint/S016")
    lines.append("")
    lines.append("## Summary")
    lines.append("")
    lines.append(f"- Reference VOR: {len(reference)} rows")
    lines.append(f"- Generated VOR: {len(generated)} rows")
    lines.append(f"- Matched rows: {stats['matched']}")
    lines.append(f"- Match rate: {stats['match_rate']:.1f}%")
    lines.append("")

    # Gate results
    lines.append("## Acceptance Gates")
    lines.append("")
    all_pass = True

    for gate_id in ["G1", "G2", "G3", "G4"]:
        gate = gate_results[gate_id]
        status = "PASS" if gate["pass"] else "FAIL"
        symbol = "✅" if gate["pass"] else "❌"

        lines.append(f"### {gate_id}: {gate['description']} {symbol}")
        lines.append("")

        if gate_id == "G1":
            lines.append(f"- **Result:** {gate['value']:.1f}%")
            lines.append(f"- **Threshold:** >= {gate['threshold']}%")
            lines.append(f"- **Status:** {status}")

        elif gate_id == "G2":
            lines.append(f"- **Reference:** {gate['ref_value']:.1f} m")
            lines.append(f"- **Generated:** {gate['gen_value']:.1f} m")
            lines.append(f"- **Target range:** {gate['min_threshold']:.0f} - {gate['max_threshold']:.0f} m")
            lines.append(f"- **Status:** {status}")

        elif gate_id == "G3":
            lines.append(f"- **Reference:** {gate['ref_value']:.0f} sht")
            lines.append(f"- **Generated:** {gate['gen_value']:.0f} sht")
            lines.append(f"- **Target range:** {gate['min_threshold']:.0f} - {gate['max_threshold']:.0f} sht")
            lines.append(f"- **Status:** {status}")

        elif gate_id == "G4":
            lines.append("| Category | Present |")
            lines.append("|----------|---------|")
            for cat, present in gate["categories"].items():
                sym = "✅" if present else "❌"
                lines.append(f"| {cat} | {sym} |")
            lines.append("")
            lines.append(f"- **Status:** {status}")

        lines.append("")
        all_pass = all_pass and gate["pass"]

    # Overall result
    lines.append("## Overall Result")
    lines.append("")
    if all_pass:
        lines.append("✅ **ALL GATES PASSED**")
    else:
        lines.append("❌ **SOME GATES FAILED**")
    lines.append("")

    # Match details
    lines.append("## Match Details")
    lines.append("")
    lines.append("### Top 10 Matched Rows")
    lines.append("")
    lines.append("| Score | Reference Name | Generated Name | Ref Qty | Gen Qty |")
    lines.append("|-------|----------------|----------------|---------|---------|")

    sorted_matches = sorted(matches, key=lambda m: m[2], reverse=True)[:10]
    for ri, gi, score in sorted_matches:
        ref_row = reference[ri]
        gen_row = generated[gi]
        ref_qty = ref_row.get("qty", 0.0)
        gen_qty = gen_row.get("total", 0.0) or gen_row.get("qty", 0.0)

        lines.append(f"| {score:.2f} | {ref_row['name'][:40]} | {gen_row['name'][:40]} | {ref_qty:.1f} | {gen_qty:.1f} |")

    lines.append("")

    # Unmatched reference rows
    lines.append("### Unmatched Reference Rows (first 20)")
    lines.append("")

    matched_ref_indices = set(m[0] for m in matches)
    unmatched_ref = [
        (i, row) for i, row in enumerate(reference)
        if i not in matched_ref_indices
    ][:20]

    if unmatched_ref:
        lines.append("| Name | Unit | Qty |")
        lines.append("|------|------|-----|")
        for i, row in unmatched_ref:
            lines.append(f"| {row['name'][:60]} | {row['unit']} | {row['qty']:.1f} |")
    else:
        lines.append("*None*")

    lines.append("")

    # Category totals
    lines.append("## Category Totals")
    lines.append("")

    ref_totals = calculate_category_totals(reference)
    gen_totals = calculate_category_totals(generated)

    all_cats = sorted(set(ref_totals.keys()) | set(gen_totals.keys()))

    lines.append("| Category | Reference | Generated |")
    lines.append("|----------|-----------|-----------|")

    for cat in all_cats:
        ref_units = ref_totals[cat]
        gen_units = gen_totals[cat]

        # Format as "X m, Y sht" etc
        ref_str = ", ".join(f"{v:.1f} {u}" for u, v in sorted(ref_units.items()))
        gen_str = ", ".join(f"{v:.1f} {u}" for u, v in sorted(gen_units.items()))

        lines.append(f"| {cat} | {ref_str or '0'} | {gen_str or '0'} |")

    lines.append("")

    # Implementation notes
    lines.append("## Implementation Notes")
    lines.append("")
    lines.append("- Used `vor_compose.compose_vor_table()` for VOR generation")
    lines.append("- Comparison algorithm: token-based name similarity with unit bonus")
    lines.append("- Match threshold: 0.6 similarity score")
    lines.append("- Category detection: keyword-based from row names")
    lines.append("")

    return "\n".join(lines)


# ============================================================================
# Main
# ============================================================================

def main():
    """Main entry point."""
    log.info("="*80)
    log.info("T074: VOR Composition Verification")
    log.info("="*80)

    # Validate paths
    if not PDF_DIR.exists():
        log.error("PDF directory not found: %s", PDF_DIR)
        return 1

    if not REFERENCE_DOCX.exists():
        log.error("Reference DOCX not found: %s", REFERENCE_DOCX)
        return 1

    # Step 1: Parse reference VOR
    log.info("\n--- Step 1: Parse reference VOR ---")
    reference = parse_reference_vor(REFERENCE_DOCX)

    if not reference:
        log.error("Failed to parse reference VOR")
        return 1

    # Step 2: Generate VOR from PDFs
    log.info("\n--- Step 2: Generate VOR from PDFs ---")
    generated = generate_vor_from_pdfs(PDF_DIR)

    if not generated:
        log.error("Failed to generate VOR")
        return 1

    # Step 3: Match and compare
    log.info("\n--- Step 3: Match and compare ---")
    matches, stats = match_vor_rows(reference, generated)

    log.info("Match statistics:")
    for key, value in stats.items():
        log.info("  %s: %s", key, value)

    # Step 4: Check acceptance gates
    log.info("\n--- Step 4: Check acceptance gates ---")
    gate_results = check_gates(reference, generated, matches, stats)

    for gate_id, gate in gate_results.items():
        status = "PASS" if gate["pass"] else "FAIL"
        log.info("%s: %s - %s", gate_id, gate["description"], status)

    # Step 5: Generate report
    log.info("\n--- Step 5: Generate report ---")
    report = generate_markdown_report(reference, generated, matches, stats, gate_results)

    # Ensure output directory exists
    OUTPUT_REPORT.parent.mkdir(parents=True, exist_ok=True)

    # Write report
    OUTPUT_REPORT.write_text(report, encoding="utf-8")
    log.info("Report saved: %s", OUTPUT_REPORT)
    log.info("Report size: %.1f KB", len(report) / 1024)

    # Summary
    log.info("\n" + "="*80)
    log.info("Verification complete!")
    log.info("="*80)

    all_pass = all(g["pass"] for g in gate_results.values())
    if all_pass:
        log.info("✅ ALL GATES PASSED")
        return 0
    else:
        log.warning("❌ SOME GATES FAILED")
        return 1


if __name__ == "__main__":
    sys.exit(main())
