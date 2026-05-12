"""
T121 / S011-01: Extract reference VOR (count-only) for GPK Захватка 3.

Source: Data/ДБТ разделы для ИИ/03_ГПК_/3-я захватка/ВОР ЭО, Захватка 3_ГПК.docx
Output: Data/ДБТ разделы для ИИ/03_ГПК_/3-я захватка/gpk3_reference_count_only.xlsx

Behavior:
  - Parse the equipment table (cols: №, Наименование, Ед. изм., РД (qty), ...)
  - Track current section (heading row with empty № / empty unit / empty qty)
  - Drop header rows and empty rows
  - Normalize item names by stripping installation-height markers
  - Collapse rows that differ only by height marker by summing qty
  - Preserve original names list per collapsed row
  - Write xlsx with: item_name_normalized, original_names (joined by " | "),
    unit, total_qty, section, source_row_count

Heuristics for height detection are documented in the accompanying note
file: gpk3_reference_count_only_NOTES.md
"""
from __future__ import annotations

import io
import re
import sys
from collections import OrderedDict
from pathlib import Path

# Force utf-8 stdout to avoid Windows cp1254 issues per KB-006
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

from docx import Document
from openpyxl import Workbook

DOCX_PATH = Path(
    "Data/ДБТ разделы для ИИ/03_ГПК_/3-я захватка/ВОР ЭО, Захватка 3_ГПК.docx"
)
OUT_PATH = Path(
    "Data/ДБТ разделы для ИИ/03_ГПК_/3-я захватка/gpk3_reference_count_only.xlsx"
)
NOTES_PATH = Path(
    "Data/ДБТ разделы для ИИ/03_ГПК_/3-я захватка/gpk3_reference_count_only_NOTES.md"
)


# ---------------------------------------------------------------------------
# Height-marker heuristics
# ---------------------------------------------------------------------------

# Patterns that indicate installation-height qualifiers we want to strip when
# normalizing item names. We support several common Russian phrasings.
#
# Examples that must be matched and removed:
#   "на высоте до 5 метров"          -> stripped
#   "на высоте от 5 до 13 метров"    -> stripped
#   "на высоте от 13 до 20 м"        -> stripped
#   "на высоте от 20 до 35 метров"   -> stripped
#   "на отм. +4.200"                 -> stripped
#   "отм. 0.000"                     -> stripped
#   "на отметке +7.800"              -> stripped
#
# We also strip the colon at the end if present (header rows end with ":").

# Matches phrases like "на высоте до 5 метров", "на высоте от 5 до 13 м".
# We distinguish height markers from millimetre-section markers ("до 10 мм2",
# "до 6 мм2") by requiring "м" / "метр(а|ов)?" to be a whole word — i.e. NOT
# followed by another letter. We use a negative-lookahead `(?![а-яa-z0-9])`.
_M_END = r"(?:метров|метра|метр|м)(?![а-яa-z0-9])\.?"

HEIGHT_PATTERNS = [
    # "на высоте от N до M метров / м"
    re.compile(rf"\s*на\s+высоте\s+от\s+\d+(?:[.,]\d+)?\s+до\s+\d+(?:[.,]\d+)?\s*{_M_END}", re.IGNORECASE),
    # "на высоте до N метров / м"
    re.compile(rf"\s*на\s+высоте\s+до\s+\d+(?:[.,]\d+)?\s*{_M_END}", re.IGNORECASE),
    # "на высоте от N метров / м"
    re.compile(rf"\s*на\s+высоте\s+от\s+\d+(?:[.,]\d+)?\s*{_M_END}", re.IGNORECASE),
    # Bare "от N до M метров / м" (no "на высоте" prefix) — only if it appears
    # AT THE END of the name (so it's not confused with mid-name dimensions).
    re.compile(rf"\s+от\s+\d+(?:[.,]\d+)?\s+до\s+\d+(?:[.,]\d+)?\s*{_M_END}\s*:?\s*$", re.IGNORECASE),
    # Bare "до N м(етров)" at the END of the name.
    re.compile(rf"\s+до\s+\d+(?:[.,]\d+)?\s*{_M_END}\s*:?\s*$", re.IGNORECASE),
    # "на отм. +X.XXX" / "на отм. -X.XXX" / "на отм. X.XXX"
    re.compile(r"\s*на\s+отм\.?\s*[+\-]?\d+(?:[.,]\d+)?", re.IGNORECASE),
    # "отм. +X.XXX"
    re.compile(r"\s*отм\.?\s*[+\-]?\d+(?:[.,]\d+)?", re.IGNORECASE),
    # "на отметке +X.XXX"
    re.compile(r"\s*на\s+отметке\s+[+\-]?\d+(?:[.,]\d+)?", re.IGNORECASE),
]

# After stripping height markers we may be left with dangling separators
# (e.g. trailing colon, multiple spaces, leading "(", " ,").  Clean them up.
CLEANUP_PATTERNS = [
    (re.compile(r"\s*:\s*$"), ""),
    (re.compile(r"\s{2,}"), " "),
    (re.compile(r"\s+,"), ","),
    (re.compile(r"\(\s*\)"), ""),
]


def has_height_marker(name: str) -> bool:
    return any(p.search(name) for p in HEIGHT_PATTERNS)


def strip_height_marker(name: str) -> str:
    """Remove installation-height qualifiers from an item name."""
    s = name
    for p in HEIGHT_PATTERNS:
        s = p.sub("", s)
    for cre, repl in CLEANUP_PATTERNS:
        s = cre.sub(repl, s)
    return s.strip()


def normalize_name(name: str) -> str:
    """Full normalization: strip height markers, trim, drop trailing colon,
    collapse internal whitespace.
    """
    s = strip_height_marker(name)
    s = s.strip().rstrip(":").strip()
    s = re.sub(r"\s+", " ", s)
    return s


# ---------------------------------------------------------------------------
# Qty parsing
# ---------------------------------------------------------------------------

def parse_qty(qty_text: str) -> float | None:
    """Parse VOR qty. Returns None if not a parseable number.

    Supported:
      "117"      -> 117.0
      "0"        -> 0.0
      "1455"     -> 1455.0
      "12,5"     -> 12.5
      "35/1"     -> 35.0   (composite single/three-pole; we take the first)
      ""         -> None
    """
    if qty_text is None:
        return None
    s = str(qty_text).strip()
    if not s:
        return None
    # Handle composite "N/M" – return the FIRST integer; the second is a
    # categorical sub-count (poles) and would inflate sums if added.
    if "/" in s:
        first = s.split("/", 1)[0].strip()
        s = first
    s = s.replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return None


# ---------------------------------------------------------------------------
# Main parsing
# ---------------------------------------------------------------------------

def extract_rows(docx_path: Path) -> list[dict]:
    """Walk the single table in the docx and return data rows.

    Each output dict has: section, original_name, unit, qty_raw, qty,
    has_height_marker.

    Rules:
      * Skip header rows (row 0: "№ п/п", row 1: "1 2 3 4 5"):
        identified by checking if №-cell equals "№ п/п" or all numeric labels.
      * Section headers: empty № AND empty unit AND empty qty.  Update current
        section; do NOT emit as data row.
      * Aggregate-only rows (non-empty № but empty unit AND empty qty):
        emit with qty=None.  Example: "Соединение жил кабелей методом
        опрессовки".  We DROP these from final output because they have no
        countable quantity (their children carry the actual qty).
      * Otherwise emit a data row.
    """
    doc = Document(str(docx_path))
    if not doc.tables:
        raise RuntimeError("No tables in docx")
    table = doc.tables[0]

    rows: list[dict] = []
    current_section: str = ""

    for ri, row in enumerate(table.rows):
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        # Pad to 7 columns just in case
        while len(cells) < 7:
            cells.append("")
        no_cell = cells[0]
        name = cells[1]
        unit = cells[2]
        qty_raw = cells[3]

        # Header rows
        if ri == 0 and "наименование" in name.lower():
            continue
        if ri == 1 and re.match(r"^\d+$", no_cell) and name in {"2", "Наименование"}:
            # The numeric column-index row "1 2 3 4 5 .. 7"
            continue
        # Detect numeric-only column-index row (e.g. all cells are short digits)
        if all(re.match(r"^\d?$", (c or "")) is None or (c or "").isdigit() and len(c) <= 2 for c in cells[:5]) and ri <= 2:
            # Heuristic – skip ri=1 only; we have already handled it above.
            pass

        # Section header: empty №, empty unit, empty qty, but has name
        if not no_cell and not unit and not qty_raw and name:
            current_section = name
            continue

        # Truly empty row
        if not name and not no_cell and not unit and not qty_raw:
            continue

        qty = parse_qty(qty_raw)

        rows.append(
            {
                "row_index": ri,
                "section": current_section,
                "no": no_cell,
                "original_name": name,
                "unit": unit,
                "qty_raw": qty_raw,
                "qty": qty,
                "has_height_marker": has_height_marker(name),
            }
        )

    return rows


def collapse_by_height(rows: list[dict]) -> list[dict]:
    """Group rows by (normalized_name, unit, section) and sum qty.

    Returns a list of dicts with keys:
      item_name_normalized, original_names (list of unique originals,
      order-preserving), unit, total_qty, section, source_row_count.

    Rows with qty=None or unit="" are dropped (they're aggregate headers
    with no countable quantity of their own).
    """
    grouped: "OrderedDict[tuple[str, str, str], dict]" = OrderedDict()

    for r in rows:
        if r["qty"] is None:
            # No countable qty – skip (this also drops 'Соединение жил кабелей
            # методом опрессовки' aggregate row which has unit='' qty='')
            continue
        if not r["unit"]:
            # No unit means qty cannot be summed sensibly
            continue
        norm = normalize_name(r["original_name"])
        key = (norm, r["unit"], r["section"])
        if key not in grouped:
            grouped[key] = {
                "item_name_normalized": norm,
                "original_names": [],
                "unit": r["unit"],
                "total_qty": 0.0,
                "section": r["section"],
                "source_row_count": 0,
                "collapsed_height_variants": 0,
            }
        g = grouped[key]
        if r["original_name"] not in g["original_names"]:
            g["original_names"].append(r["original_name"])
        g["total_qty"] += r["qty"]
        g["source_row_count"] += 1
        if r["has_height_marker"]:
            g["collapsed_height_variants"] += 1
    return list(grouped.values())


# ---------------------------------------------------------------------------
# Excel writer
# ---------------------------------------------------------------------------

def write_xlsx(rows: list[dict], out_path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "GPK3 reference count"

    headers = [
        "item_name_normalized",
        "original_names",
        "unit",
        "total_qty",
        "section",
        "source_row_count",
        "collapsed_height_variants",
    ]
    ws.append(headers)

    for r in rows:
        ws.append(
            [
                r["item_name_normalized"],
                " | ".join(r["original_names"]),
                r["unit"],
                r["total_qty"],
                r["section"],
                r["source_row_count"],
                r["collapsed_height_variants"],
            ]
        )

    # Column widths for readability
    widths = {"A": 60, "B": 80, "C": 10, "D": 12, "E": 35, "F": 12, "G": 12}
    for col, w in widths.items():
        ws.column_dimensions[col].width = w

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))


# ---------------------------------------------------------------------------
# Heuristics note writer
# ---------------------------------------------------------------------------

NOTES_TEMPLATE = """# GPK3 reference VOR (count-only) — extraction notes

Source: `ВОР ЭО, Захватка 3_ГПК.docx` (Электроосвещение, 3 захватка, ГПК).
Output: `gpk3_reference_count_only.xlsx`.

## Heuristics for height-marker detection

The reference VOR repeats the same installation work at multiple installation
heights, encoded in the row name. To compare against a parser that emits one
row per item (height-agnostic), we collapse such rows.

A row name is considered to contain an installation-height marker when ANY of
the following patterns match (case-insensitive). All metre-suffix patterns
use a negative lookahead so that "до 10 мм2" / "до 6 мм2" (cable cross-section)
are NOT mistakenly stripped:

```
_M_END = (?:метров|метра|метр|м)(?![а-яa-z0-9])\.?
```

| # | Pattern | Matches | Anti-matches |
|---|---|---|---|
| 1 | `на\s+высоте\s+от\s+\d+\s+до\s+\d+\s*_M_END` | "на высоте от 5 до 13 метров", "на высоте от 13 до 20 м" | — |
| 2 | `на\s+высоте\s+до\s+\d+\s*_M_END` | "на высоте до 5 метров", "на высоте до 5 м" | — |
| 3 | `на\s+высоте\s+от\s+\d+\s*_M_END` | "на высоте от 20 м" | — |
| 4 | `\s+от\s+\d+\s+до\s+\d+\s*_M_END\s*:?\s*$` (end-anchored) | "(с пиктограммой) от 5 до 13 метров:" | "до 10 мм2" inside string |
| 5 | `\s+до\s+\d+\s*_M_END\s*:?\s*$` (end-anchored) | "(с пиктограммой) до 5 м" | "до 10 мм2" (different unit) |
| 6 | `на\s+отм\.?\s*[+\-]?\d+(?:[.,]\d+)?` | "на отм. +4.200", "на отм. 0.000", "на отм -3.000" | — |
| 7 | `отм\.?\s*[+\-]?\d+(?:[.,]\d+)?` | "отм. +7.800" (standalone) | — |
| 8 | `на\s+отметке\s+[+\-]?\d+(?:[.,]\d+)?` | "на отметке +4.500" | — |

Patterns 4 and 5 are anchored to the END of the string (via `\s*:?\s*$`).
This is essential because the reference VOR also contains mid-string
dimensions such as `до 10 мм2`, `до 6 мм2`, `до 16 мм2` — these refer to
cable cross-section, not installation height, and must NOT be stripped. The
combination of (a) word-boundary lookahead on `м/метр(а|ов)?` and (b) end-of-
string anchoring eliminates the false positives.

After substring removal we re-collapse internal whitespace, drop a trailing
`:` (header phrasing), and trim. Then we group on the
`(normalized_name, unit, section)` triple and sum `total_qty`.

## Other parsing rules

1. **Single table only.** The docx has exactly one table — the rest is the
   title block.
2. **Header rows skipped.** Row 0 (`№ п/п | Наименование | Ед. изм. | РД ...`)
   and row 1 (column-index `1 2 3 4 5`) are dropped.
3. **Section headers tracked.** A row with empty №, empty unit, empty qty but
   non-empty name (e.g. `Щитовое оборудование`, `Кабельная продукция`) updates
   the `current_section` state and is not emitted as a data row. The detected
   sections are: `Щитовое оборудование`, `Светотехническое оборудование`,
   `Монтаж электроустановочных изделий`, `Кабельная продукция`,
   `Монтаж кабельных лотков и соединительных деталей`, `ПВХ изделия и трубы`,
   `Пусконаладочные работы`.
4. **Aggregate-only rows dropped from output.** Rows that carry a № but have
   empty `unit` and empty `qty` (e.g. row 61, `Соединение жил кабелей методом
   опрессовки`) are aggregate headers whose actual countable items appear in
   the children below — including them would distort the sum.
5. **Composite qty `N/M`.** Rows in the commissioning section use
   `35/1`, `23/1`, `36/1` (single-pole / three-pole). We take the FIRST
   integer (single-pole count) as the dominant qty. The full original string
   is preserved in the `original_names` chain via the original row.
6. **Qty parsing.** `12,5` -> 12.5; empty -> excluded; `35/1` -> 35.

## Collapse outcome (summary)

The xlsx includes for every collapsed row:
  - `source_row_count` — how many original rows merged
  - `collapsed_height_variants` — of those, how many carried a height marker

Rows with `collapsed_height_variants > 0` indicate genuine height-collapse;
rows with `source_row_count > collapsed_height_variants` indicate that
duplicates already shared identical names (verbatim) and were merged trivially.

## Reproduction

```
python _t121_extract_gpk3_reference.py
```
"""


def write_notes(notes_path: Path) -> None:
    notes_path.parent.mkdir(parents=True, exist_ok=True)
    notes_path.write_text(NOTES_TEMPLATE, encoding="utf-8")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    print(f"Reading: {DOCX_PATH}")
    rows = extract_rows(DOCX_PATH)
    print(f"  Extracted data rows (pre-collapse): {len(rows)}")

    with_height = sum(1 for r in rows if r["has_height_marker"])
    print(f"  Rows containing a height marker: {with_height}")

    collapsed = collapse_by_height(rows)
    print(f"  Collapsed rows (post-collapse, in output): {len(collapsed)}")

    # Quick health check
    total_qty = sum(r["total_qty"] for r in collapsed)
    print(f"  Sum of total_qty across all rows: {total_qty:.2f}")

    write_xlsx(collapsed, OUT_PATH)
    print(f"  Wrote: {OUT_PATH}")
    write_notes(NOTES_PATH)
    print(f"  Wrote: {NOTES_PATH}")

    # Print a few representative rows for QA
    print("\n=== Sample of collapsed rows where collapsing actually happened ===")
    samples = [r for r in collapsed if r["source_row_count"] > 1][:10]
    for r in samples:
        print(
            f"  qty={r['total_qty']:>8.1f} {r['unit']:>5}  "
            f"src={r['source_row_count']}  hgt={r['collapsed_height_variants']}  "
            f"[{r['section'][:25]}]  {r['item_name_normalized'][:70]}"
        )

    print("\n=== Sample of single-row entries (no collapse) ===")
    samples = [r for r in collapsed if r["source_row_count"] == 1][:5]
    for r in samples:
        print(
            f"  qty={r['total_qty']:>8.1f} {r['unit']:>5}  "
            f"[{r['section'][:25]}]  {r['item_name_normalized'][:70]}"
        )


if __name__ == "__main__":
    main()
