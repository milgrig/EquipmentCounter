import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

PAIRS = [
    ('ETALON', r'Data\ДБТ разделы для ИИ\03_ГПК_\3-я захватка\ВОР ЭО, Захватка 3_ГПК.docx'),
    ('CURRENT', r'Data\ДБТ разделы для ИИ\03_ГПК_\3-я захватка\VOR_02_PDF.docx'),
]

for label, p in PAIRS:
    print('=' * 80)
    print(label, p)
    print('=' * 80)
    d = Document(p)
    print('Paragraphs:', len(d.paragraphs), 'Tables:', len(d.tables))
    print()
    print('--- PARAGRAPHS ---')
    for i, par in enumerate(d.paragraphs):
        t = par.text.strip()
        if t:
            sty = par.style.name
            print('  [%d] style=%r | %s' % (i, sty, t[:140]))

    for ti, tbl in enumerate(d.tables):
        print()
        print('--- TABLE %d: rows=%d cols=%d ---' % (ti, len(tbl.rows), len(tbl.columns)))
        for ri, row in enumerate(tbl.rows):
            cells = []
            for c in row.cells:
                txt = c.text.strip().replace('\n', ' | ')
                cells.append(txt[:50])
            print('  R%d: %s' % (ri, cells))
            if ri >= 12 and ri < len(tbl.rows) - 3:
                print('  ... skipping middle')
                # jump to last 3
                for ri2 in range(len(tbl.rows) - 3, len(tbl.rows)):
                    row2 = tbl.rows[ri2]
                    cells2 = [c.text.strip().replace('\n', ' | ')[:50] for c in row2.cells]
                    print('  R%d: %s' % (ri2, cells2))
                break
    print()
