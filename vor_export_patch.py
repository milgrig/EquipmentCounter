"""
vor_export_patch.py — расширяет web_app.app:
- кэш агрегата ВОР по folder_id (TTL 30 мин)
- быстрый /api/folder/{folder_id}/export_xlsx_v2 (из кэша)
- новый      /api/folder/{folder_id}/export_docx
- хук для записи кэша при завершении SSE-обработки

Подключение (одной строкой в конце web_app.py):
    from vor_export_patch import register_vor_export
    register_vor_export(app)
"""

from __future__ import annotations

import asyncio
import io
import time
from pathlib import Path

from fastapi import HTTPException
from fastapi.responses import Response


_CACHE_TTL = 1800.0  # 30 минут


def _get_cache(app):
    if not hasattr(app.state, "vor_cache"):
        app.state.vor_cache = {}
    return app.state.vor_cache


def store_vor_result(app, folder_id: str, *, aggregated, all_results, rel_folder):
    """Положить агрегированный ВОР в кэш — вызывается из SSE-обработчика."""
    _get_cache(app)[folder_id] = {
        "aggregated": aggregated,
        "all_results": all_results,
        "rel_folder": rel_folder,
        "ts": time.time(),
    }


def get_cached_vor(app, folder_id: str):
    cache = _get_cache(app)
    entry = cache.get(folder_id)
    if not entry:
        return None
    if time.time() - entry["ts"] > _CACHE_TTL:
        cache.pop(folder_id, None)
        return None
    return entry


def _build_docx_bytes(aggregated, rel_folder: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    h = doc.add_heading("Ведомость объёмов работ", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(f"Папка: {rel_folder}")
    run.bold = True

    headers = ["№", "Наименование вида работ", "Ед.", "Объём",
               "Формула расчёта", "Ссылка на чертежи"]
    widths_cm = [1.0, 8.0, 1.5, 2.0, 3.5, 3.5]

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, (txt, w) in enumerate(zip(headers, widths_cm)):
        hdr[i].text = txt
        hdr[i].width = Cm(w)
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    for row in aggregated:
        cells = tbl.add_row().cells
        cells[0].text = str(row.get("row", ""))
        cells[1].text = str(row.get("name", ""))
        cells[2].text = str(row.get("unit", ""))
        cells[3].text = str(row.get("total", ""))
        cells[4].text = str(row.get("formula", ""))
        cells[5].text = str(row.get("drawing_refs", ""))
        for c, w in zip(cells, widths_cm):
            c.width = Cm(w)
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def register_vor_export(app):
    """Регистрирует на FastAPI-приложении кэш-эндпоинты XLSX/DOCX."""
    # Берём приватные хелперы из web_app, чтобы fallback-пересчёт работал
    import web_app as wa

    @app.get("/api/folder/{folder_id}/export_docx")
    async def api_folder_export_docx(folder_id: str):
        rel_folder = wa._id_to_folder(folder_id)
        if rel_folder is None:
            raise HTTPException(status_code=404, detail="Folder not found")

        cached = get_cached_vor(app, folder_id)
        if cached:
            aggregated = cached["aggregated"]
        else:
            pdf_files = wa._folder_files(rel_folder)
            if not pdf_files:
                raise HTTPException(status_code=404, detail="No PDFs in folder")
            all_results: dict[str, list[dict]] = {}
            for pdf_path in pdf_files:
                try:
                    file_result = await asyncio.to_thread(
                        wa._count_equipment_in_pdf, str(pdf_path)
                    )
                    all_results[pdf_path.name] = file_result
                except Exception:
                    continue
            aggregated = wa._aggregate_equipment(all_results)
            store_vor_result(
                app, folder_id,
                aggregated=aggregated,
                all_results=all_results,
                rel_folder=rel_folder,
            )

        content = await asyncio.to_thread(_build_docx_bytes, aggregated, rel_folder)

        folder_name = rel_folder.rsplit("/", 1)[-1] if "/" in rel_folder else rel_folder
        # B044 KB-006: HTTP headers must be latin-1; use ASCII fallback +
        # RFC 5987 filename* for UTF-8 (browsers prefer filename* when present).
        ascii_safe = "".join(
            c if (c.isascii() and (c.isalnum() or c in "._- ")) else "_"
            for c in folder_name
        ) or "folder"
        from urllib.parse import quote as _q
        utf8_q = _q(f"VOR_{folder_name}.docx", safe="")
        return Response(
            content=content,
            media_type=("application/vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"),
            headers={
                "Content-Disposition": (
                    f'attachment; filename="VOR_{ascii_safe}.docx"; '
                    f"filename*=UTF-8''{utf8_q}"
                ),
            },
        )

    @app.get("/api/folder/{folder_id}/export_xlsx_v2")
    async def api_folder_export_xlsx_v2(folder_id: str):
        """Быстрый XLSX из кэша. Если кэша нет — пересчёт и кэширование."""
        import openpyxl
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

        rel_folder = wa._id_to_folder(folder_id)
        if rel_folder is None:
            raise HTTPException(status_code=404, detail="Folder not found")

        cached = get_cached_vor(app, folder_id)
        if cached:
            aggregated = cached["aggregated"]
        else:
            pdf_files = wa._folder_files(rel_folder)
            if not pdf_files:
                raise HTTPException(status_code=404, detail="No PDFs in folder")
            all_results: dict[str, list[dict]] = {}
            for pdf_path in pdf_files:
                try:
                    file_result = await asyncio.to_thread(
                        wa._count_equipment_in_pdf, str(pdf_path)
                    )
                    all_results[pdf_path.name] = file_result
                except Exception:
                    continue
            aggregated = wa._aggregate_equipment(all_results)
            store_vor_result(
                app, folder_id,
                aggregated=aggregated,
                all_results=all_results,
                rel_folder=rel_folder,
            )

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "ВОР"

        header_font = Font(bold=True, size=10)
        header_fill = PatternFill("solid", fgColor="D9E1F2")
        thin = Border(left=Side("thin"), right=Side("thin"),
                      top=Side("thin"), bottom=Side("thin"))

        headers = ["№ п/п", "Наименование вида работ", "Ед. изм.", "Объем работ",
                   "Формула расчета", "Ссылка на чертежи", "Доп. информация"]
        col_widths = [7, 72, 9, 12, 20, 26, 27]
        for ci, (h, w) in enumerate(zip(headers, col_widths), 1):
            cell = ws.cell(row=1, column=ci, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin
            cell.alignment = Alignment(horizontal="center", wrap_text=True)
            ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = w

        for row_data in aggregated:
            ri = row_data["row"] + 1
            vals = [row_data["row"], row_data["name"], row_data["unit"],
                    row_data["total"], row_data["formula"], row_data["drawing_refs"],
                    row_data.get("extra_info", "")]
            for ci, v in enumerate(vals, 1):
                cell = ws.cell(row=ri, column=ci, value=v)
                cell.border = thin
                if ci in (1, 3, 4):
                    cell.alignment = Alignment(horizontal="center")

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)

        folder_name = rel_folder.rsplit("/", 1)[-1] if "/" in rel_folder else rel_folder
        # B044 KB-006: HTTP headers latin-1 only -> RFC 5987 filename*
        ascii_safe = "".join(
            c if (c.isascii() and (c.isalnum() or c in "._- ")) else "_"
            for c in folder_name
        ) or "folder"
        from urllib.parse import quote as _q
        utf8_q = _q(f"VOR_{folder_name}.xlsx", safe="")
        return Response(
            content=buf.getvalue(),
            media_type=("application/vnd.openxmlformats-officedocument."
                        "spreadsheetml.sheet"),
            headers={
                "Content-Disposition": (
                    f'attachment; filename="VOR_{ascii_safe}.xlsx"; '
                    f"filename*=UTF-8''{utf8_q}"
                ),
            },
        )

    @app.post("/api/folder/{folder_id}/cache_vor")
    async def api_folder_cache_vor(folder_id: str, payload: dict):
        """Принять уже посчитанный агрегат от клиента (после SSE done)
        и положить его в кэш для последующих /export_xlsx_v2 и /export_docx.
        """
        rel_folder = wa._id_to_folder(folder_id)
        if rel_folder is None:
            raise HTTPException(status_code=404, detail="Folder not found")
        aggregated = payload.get("aggregated") or []
        store_vor_result(
            app, folder_id,
            aggregated=aggregated,
            all_results={},
            rel_folder=rel_folder,
        )
        return {"status": "ok", "rows": len(aggregated)}

    return app
