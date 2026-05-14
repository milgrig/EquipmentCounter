"""T087 smoke test — render a DOCX and verify height section headers.

Builds a synthetic aggregated list (shape matches what
`vor_elevation_grouper.regroup_aggregate_by_elevation` produces) with
luminaire rows whose `per_elevation` spans all four height buckets, plus
rows in other sections.  Renders via `render_vor_docx`, opens the result
with python-docx, and asserts that the table contains the four expected
bold sub-header rows produced by T087 grouping.
"""
from __future__ import annotations

import io
from pathlib import Path

from docx import Document

from vor_docx_renderer import (
    HEIGHT_BUCKET_PHRASES,
    HEIGHT_BUCKET_ORDER,
    _derive_height_bucket_from_elevations,
    _elevation_to_bucket,
    _group_luminaires_by_height,
    _resolve_height_bucket,
    render_vor_docx,
)

# Unit checks on the helpers first.
assert _elevation_to_bucket("0.000") == "<5m"
assert _elevation_to_bucket("+4.200") == "<5m"
assert _elevation_to_bucket("+7.800") == "5-13m"
assert _elevation_to_bucket("+13.800") == "13-20m"
assert _elevation_to_bucket("+18.600") == "13-20m"
assert _elevation_to_bucket("+23.400") == "20-35m"
assert _elevation_to_bucket("+28.200") == "20-35m"
assert _elevation_to_bucket("") == "unknown"
assert _elevation_to_bucket(None) == "unknown"
print("OK: _elevation_to_bucket")

assert _derive_height_bucket_from_elevations({"0.000": 10}) == "<5m"
assert _derive_height_bucket_from_elevations(
    {"0.000": 1, "+4.200": 1, "+7.800": 99}
) == "5-13m"
assert _derive_height_bucket_from_elevations({}) == "unknown"
assert _derive_height_bucket_from_elevations(None) == "unknown"
print("OK: _derive_height_bucket_from_elevations")

assert _resolve_height_bucket({"_height_bucket": "13-20m"}) == "13-20m"
assert _resolve_height_bucket({"height_bucket": "20-35m"}) == "20-35m"
assert _resolve_height_bucket(
    {"per_elevation": {"+18.600": 4}}
) == "13-20m"
assert _resolve_height_bucket({}) == "unknown"
print("OK: _resolve_height_bucket")

# Build aggregated input.
aggregated = [
    # Щитовое
    {"name": "Монтаж щита ЩО-3", "unit": "шт", "total": 1, "formula": "1",
     "drawing_refs": "003 - Схемы ЩО, ЩАО-ЩО-3.pdf"},
    # Светотехническое — 4 bucket-а
    {"name": "Светильник ARCTIC.OPL ECO LED 1200 5000K",
     "unit": "шт", "total": 14, "formula": "14",
     "drawing_refs": "005-Планы освещения-отм. 0.000.pdf",
     "per_elevation": {"0.000": 14}},
    {"name": "Светильник SLICK.PRS LED 30 with driver box Ex 5000K",
     "unit": "шт", "total": 8, "formula": "8",
     "drawing_refs": "007-Планы освещения-отм. +7.800 +9.000.pdf",
     "per_elevation": {"+7.800": 8}},
    {"name": "Светильник INSEL LB/S LED 120 D60 Ex 5000K",
     "unit": "шт", "total": 5, "formula": "5",
     "drawing_refs": "009-Планы освещения-отм. +18.600.pdf",
     "per_elevation": {"+18.600": 5}},
    {"name": "Светильник INSEL LB/S LED 120 D60 Ex 5000K (верх)",
     "unit": "шт", "total": 4, "formula": "4",
     "drawing_refs": "011-Планы освещения-отм. +28.200.pdf",
     "per_elevation": {"+28.200": 4}},
    # Розетки
    {"name": "Розетка скрытая 16А", "unit": "шт", "total": 12, "formula": "12",
     "drawing_refs": "012-План привязки на отм. 0.000.pdf"},
    # Кабельная продукция
    {"name": "Кабель ВБШвнг(А)-LS 3х1,5", "unit": "м", "total": 503,
     "formula": "503",
     "drawing_refs": "СО-СО_02.pdf"},
]

bytes_ = render_vor_docx(
    aggregated,
    rel_folder="ЭО, Захватка 3_ГПК",
    issue_date="14.05.2026",
)
print(f"OK: render_vor_docx returned {len(bytes_)} bytes")

# Re-open and inspect.
doc = Document(io.BytesIO(bytes_))
assert len(doc.tables) == 1, f"expected 1 table, got {len(doc.tables)}"
tbl = doc.tables[0]

# Collect col-2 (Наименование) cell text for every row, plus whether
# the cell's first run is bold.
rows_info: list[tuple[str, bool]] = []
for r in tbl.rows:
    cell = r.cells[1]
    txt = cell.text.strip()
    bold = False
    if cell.paragraphs:
        for p in cell.paragraphs:
            for run in p.runs:
                if run.bold and run.text.strip():
                    bold = True
                    break
            if bold:
                break
    rows_info.append((txt, bold))

# Print for diagnostic.
print("\n=== Table row 2nd-column dump (text, bold) ===")
for i, (txt, bold) in enumerate(rows_info):
    marker = "[B]" if bold else "   "
    print(f"  row {i:2d} {marker} {txt!r}")

# Assertions.
bold_rows = [t for t, b in rows_info if b]
required_subheaders = [
    "Монтаж светильников на шпильках к перекрытию на высоте до 5 метров",
    "Монтаж светильников на шпильках к перекрытию на высоте от 5 до 13 метров",
    "Монтаж светильников на шпильках к перекрытию на высоте от 13 до 20 метров",
    "Монтаж светильников на шпильках к перекрытию на высоте от 20 до 35 метров",
]
missing = [h for h in required_subheaders if h not in bold_rows]
assert not missing, f"missing height sub-headers in bold rows: {missing}"
print(f"\nOK: all 4 height sub-headers present as bold rows")

# Verify section header «Светотехническое оборудование» still present.
assert "Светотехническое оборудование" in bold_rows, (
    "main section header lost"
)
print("OK: parent section header 'Светотехническое оборудование' present")

# Verify ordering: главная секция должна быть перед всеми 4 sub-headers.
parent_idx = bold_rows.index("Светотехническое оборудование")
for h in required_subheaders:
    sub_idx = bold_rows.index(h)
    assert sub_idx > parent_idx, f"sub-header {h!r} appears before parent"
print("OK: 4 sub-headers come after parent section header")

# Verify the 4 sub-headers are in canonical ascending-height order.
sub_indices = [bold_rows.index(h) for h in required_subheaders]
assert sub_indices == sorted(sub_indices), "sub-headers out of order"
print("OK: 4 sub-headers in ascending-height canonical order")

# Save the docx for boss to open in Word.
out = Path("_t087_sample.docx")
out.write_bytes(bytes_)
print(f"\nWrote sample: {out.resolve()}")
print("\n=== T087 smoke PASSED ===")
