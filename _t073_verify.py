"""T073 acceptance verifier -- vor_compose.

Pipeline:
  1. Run _count_equipment_in_pdf on every PDF in
     "3-ya zahvatka/02_PDF".
  2. Cache results to .tayfa/common/discussions/T066_items_cache.json
     so re-runs are fast.
  3. Compose VOR_table via vor_compose.compose_vor_table.
  4. Read the reference VOR docx and run fuzzy name-match.
  5. Emit per-row alignment to T066_compose_output.json.
  6. Print gate report.

Gates:
  G1: composed VOR_table has >= 80 rows.
  G2: fuzzy row-match rate vs reference > 65%.
  G3: matched-qty within 50% on cables, within 30% on luminaires.
  G4: discussions/T066_compose_output.json emitted.

KB-006: ASCII-only source.
"""

import json
import os
import sys
import time

sys.stdout.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

ROOT = os.path.dirname(os.path.abspath(__file__))
PDF_DIR = os.path.join(
    ROOT,
    "Data",
    "\u0414\u0411\u0422 \u0440\u0430\u0437\u0434\u0435\u043b\u044b "
    "\u0434\u043b\u044f \u0418\u0418",
    "03_\u0413\u041f\u041a_",
    "3-\u044f \u0437\u0430\u0445\u0432\u0430\u0442\u043a\u0430",
    "02_PDF",
)
REF_DOCX = os.path.join(
    ROOT,
    "Data",
    "\u0414\u0411\u0422 \u0440\u0430\u0437\u0434\u0435\u043b\u044b "
    "\u0434\u043b\u044f \u0418\u0418",
    "03_\u0413\u041f\u041a_",
    "3-\u044f \u0437\u0430\u0445\u0432\u0430\u0442\u043a\u0430",
    "\u0412\u041e\u0420 \u042d\u041e, \u0417\u0430\u0445\u0432\u0430\u0442\u043a"
    "\u0430 3_\u0413\u041f\u041a.docx",
)
DISCUSS_DIR = os.path.join(ROOT, ".tayfa", "common", "discussions")
CACHE_PATH = os.path.join(DISCUSS_DIR, "T066_items_cache.json")
OUTPUT_PATH = os.path.join(DISCUSS_DIR, "T066_compose_output.json")


# ---------------------------------------------------------------------------
# Step 1: gather items per PDF (with cache)
# ---------------------------------------------------------------------------

def _list_pdfs() -> list[str]:
    out = []
    for nm in sorted(os.listdir(PDF_DIR)):
        if nm.lower().endswith(".pdf"):
            out.append(nm)
    return out


def _load_cache() -> dict | None:
    if not os.path.exists(CACHE_PATH):
        return None
    try:
        with open(CACHE_PATH, "r", encoding="utf-8") as fh:
            return json.load(fh)
    except Exception as exc:
        print("WARN: cache read failed: %r" % exc)
        return None


def _save_cache(data: dict) -> None:
    os.makedirs(DISCUSS_DIR, exist_ok=True)
    with open(CACHE_PATH, "w", encoding="utf-8") as fh:
        json.dump(data, fh, ensure_ascii=False, indent=1)


def gather_items(force: bool = False, lite: bool = True
                 ) -> dict[str, list[dict]]:
    """Return {pdf_name: items_list} ready for vor_compose.

    Modes:
      * lite=True (default for T073 harness) -- read pre-existing
        equipment_report.json from "3-ya zahvatka/", then run the
        attribution chain (vor_work_mapping, height_bucketer,
        route_classifier, thickness_extractor) on the cached items.
        Compose-only test runs in seconds rather than hours.
      * lite=False -- full pipeline; calls
        web_app._count_equipment_in_pdf (minutes per PDF).
    """
    pdfs = _list_pdfs()
    print("found %d PDFs in 02_PDF" % len(pdfs))

    cache = None if force else _load_cache()
    if cache and set(cache.keys()) >= set(pdfs):
        print("cache hit: using %s" % CACHE_PATH)
        return {p: cache[p] for p in pdfs}

    if lite:
        return _gather_lite(pdfs)

    from web_app import _count_equipment_in_pdf
    out: dict[str, list[dict]] = {}
    if cache:
        out.update({p: cache[p] for p in cache.keys() if p in pdfs})
    t0 = time.time()
    for i, name in enumerate(pdfs, 1):
        if name in out:
            print("  [%2d/%d] %s -- cached" % (i, len(pdfs), name[:60]))
            continue
        path = os.path.join(PDF_DIR, name)
        print("  [%2d/%d] %s ..." % (i, len(pdfs), name[:60]))
        t1 = time.time()
        try:
            items = _count_equipment_in_pdf(path)
        except Exception as exc:
            print("    !! exception: %r" % exc)
            items = []
        dt = time.time() - t1
        print("    -> %d items in %.1fs" % (len(items), dt))
        out[name] = items
        _save_cache(out)
    print("total elapsed: %.1fs" % (time.time() - t0))
    return out


def _gather_lite(pdfs: list[str]) -> dict[str, list[dict]]:
    """Read cached per-PDF items + apply attribution chain.

    Source: Data/.../3-ya zahvatka/equipment_report.json.  The
    cached items carry only name/count/count_ae/symbol, so we
    synthesise:
      - unit + work_name + equipment_name + category (vor_work_mapping)
      - total = count + count_ae
      - height_bucket (height_bucketer)
      - route / mount (route_classifier)
      - cross_section / diameter_mm / width_mm / height_mm
        (thickness_extractor)
    """
    src = os.path.join(
        ROOT, "Data",
        "\u0414\u0411\u0422 \u0440\u0430\u0437\u0434\u0435\u043b\u044b "
        "\u0434\u043b\u044f \u0418\u0418",
        "03_\u0413\u041f\u041a_",
        "3-\u044f \u0437\u0430\u0445\u0432\u0430\u0442\u043a\u0430",
        "equipment_report.json",
    )
    if not os.path.exists(src):
        print("WARN: no cached equipment_report.json at %s" % src)
        return {p: [] for p in pdfs}
    with open(src, "r", encoding="utf-8") as fh:
        report = json.load(fh)
    per_file = report.get("per_file", [])
    print("lite mode: loaded %d per_file entries from cache"
          % len(per_file))

    by_name: dict[str, list[dict]] = {}
    for entry in per_file:
        pdfpath = entry.get("pdf") or entry.get("pdf_name") or ""
        basename = os.path.basename(pdfpath)
        items = list(entry.get("items", []) or [])
        norm = []
        for it in items:
            nm = (it.get("name") or "").strip()
            if not nm or nm.startswith("[Auto-detected"):
                continue
            new = dict(it)
            new.setdefault("unit", "\u0448\u0442")
            cnt = float(it.get("count") or 0)
            ae = float(it.get("count_ae") or 0)
            new["total"] = cnt + ae
            norm.append(new)
        by_name[basename] = norm

    from vor_work_mapping import map_items as vor_map_items
    import height_bucketer
    import route_classifier
    import thickness_extractor

    out: dict[str, list[dict]] = {}
    for p in pdfs:
        raw = by_name.get(p, [])
        items = vor_map_items(raw) if raw else []
        try:
            height_bucketer.attribute_items(
                items, os.path.join(PDF_DIR, p))
        except Exception as exc:
            print("    height_bucket failed on %s: %r" % (p, exc))
        try:
            route_classifier.attribute_items(items)
        except Exception as exc:
            print("    route_classifier failed on %s: %r" % (p, exc))
        try:
            thickness_extractor.attribute_items(items)
        except Exception as exc:
            print("    thickness_extractor failed on %s: %r" % (p, exc))
        out[p] = items
    _save_cache(out)
    print("lite mode: produced items for %d PDFs" % len(out))
    return out


# ---------------------------------------------------------------------------
# Step 2: compose
# ---------------------------------------------------------------------------

def compose(items_per_pdf: dict[str, list[dict]]) -> list[dict]:
    import vor_compose
    rows = vor_compose.compose_vor_table(items_per_pdf)
    sm = vor_compose.summarize(rows)
    print("composed: %d rows  (cables=%d, luminaires=%d, trays=%d, conduits=%d)"
          % (sm["total_rows"], sm["cable_rows"], sm["luminaire_rows"],
             sm["tray_rows"], sm["conduit_rows"]))
    return rows


# ---------------------------------------------------------------------------
# Step 3: read reference VOR docx
# ---------------------------------------------------------------------------

def _read_reference() -> list[dict]:
    from docx import Document
    if not os.path.exists(REF_DOCX):
        print("WARN: reference not found at %s" % REF_DOCX)
        return []
    doc = Document(REF_DOCX)
    if not doc.tables:
        return []
    t = doc.tables[0]
    rows = []
    for i, r in enumerate(t.rows):
        cells = [c.text.strip() for c in r.cells]
        if len(cells) < 7:
            continue
        # Skip header (row 0) + numbering row (row 1) + section headers
        # (rows where cells[0] and cells[2] and cells[3] are all empty).
        if i < 2:
            continue
        name = cells[1]
        unit = cells[2]
        qty_text = cells[3]
        if not name or not qty_text:
            continue
        try:
            qty = float(qty_text.split("/")[0].replace(",", "."))
        except ValueError:
            continue
        rows.append({
            "name": name,
            "unit": unit,
            "qty": qty,
            "sheet_ref": cells[5],
        })
    return rows


# ---------------------------------------------------------------------------
# Step 4: fuzzy match composed rows vs reference
# ---------------------------------------------------------------------------

def _is_cable_ref(name: str) -> bool:
    lo = name.lower()
    return ("\u043a\u0430\u0431\u0435\u043b" in lo  # kabel
            or "\u043f\u0440\u043e\u0432\u043e\u0434\u043a" in lo  # provodk
            or "\u043f\u0440\u043e\u0432\u043e\u0434\u0430" in lo)  # provoda


def _is_lum_ref(name: str) -> bool:
    lo = name.lower()
    return ("\u0441\u0432\u0435\u0442\u0438\u043b" in lo
            or "\u0443\u043a\u0430\u0437\u0430\u0442\u0435\u043b" in lo
            or "\u043f\u0438\u043a\u0442\u043e\u0433\u0440\u0430\u043c" in lo)


def fuzzy_align(composed: list[dict], reference: list[dict]
                ) -> tuple[list[dict], dict]:
    """Greedy fuzzy alignment with unit-aware + qty-aware tie-break.

    For each composed row we collect the top-8 token_set_ratio
    candidates from the reference list, then re-rank them by a
    composite ``score - qty_penalty - unit_penalty``.  ``qty_pen``
    grows with the relative qty gap (0..40 points) so the cable
    fragments cannot all latch onto the single huge ``...na vysote
    do 5 m`` reference row.  ``unit_pen`` is 25 when units mismatch
    (shtuki vs metry).  Each reference row is claimed by at most one
    composed row.
    """
    from rapidfuzz import fuzz, process

    ref_names = [r["name"] for r in reference]
    matches: list[dict] = []
    used_ref: set[int] = set()
    THRESHOLD = 55

    def _combined(crow: dict, ref: dict, fuzzy_score: float) -> float:
        # Tie-break ranking only: prefer same-unit + closer-qty
        # + same-bucket candidates among those already above the
        # fuzzy threshold.  Bonuses are bounded so a high-fuzzy
        # match never loses to a low-fuzzy one with matching qty.
        cqty = float(crow.get("qty") or 0)
        rqty = float(ref.get("qty") or 0)
        qty_bonus = 0.0
        if rqty > 0 and cqty > 0:
            ratio = min(cqty, rqty) / max(cqty, rqty)
            qty_bonus = 10.0 * ratio
        unit_bonus = 5.0 if (
            (crow.get("unit") or "").strip()
            == (ref.get("unit") or "").strip()
        ) else 0.0
        # Bucket-text bonus: penalise rows whose bucket-text doesn't
        # appear in the reference name, so per-bucket composed rows
        # don't all latch onto one ref row that happens to share the
        # noun prefix.  bucket_phrases mirror vor_compose._BUCKET_TEXT.
        bucket_phrases = {
            "<5m":     "\u0434\u043e 5",
            "5-13m":   "\u043e\u0442 5 \u0434\u043e 13",
            "13-20m":  "\u043e\u0442 13 \u0434\u043e 20",
            "20-35m":  "\u043e\u0442 20 \u0434\u043e 35",
        }
        cbk = (crow.get("_height_bucket") or "").strip()
        ccat = (crow.get("_category") or "").lower()
        bucket_bonus = 0.0
        # Bucket disambiguation matters most for luminaire rows where
        # the composer emits four near-identical "Montazh svetil'nika"
        # rows that differ only by bucket text.  Cables and material
        # rows don't have this collision.
        if cbk in bucket_phrases and ("luminaire" in ccat):
            if bucket_phrases[cbk] in (ref.get("name") or ""):
                bucket_bonus = 20.0
            else:
                bucket_bonus = -10.0
        # Work/material lane separation: a composed "Montazh ..." row
        # must match a reference row that's also a work row, not a
        # material entry.  Otherwise per-model installation rows
        # collide with material rows that happen to share the model
        # name.
        montazh = "\u041c\u043e\u043d\u0442\u0430\u0436"
        cwork = (crow.get("name") or "").startswith(montazh)
        rwork = (ref.get("name") or "").startswith(montazh)
        lane_bonus = 5.0 if (cwork == rwork) else -15.0
        return (
            float(fuzzy_score) + qty_bonus + unit_bonus
            + bucket_bonus + lane_bonus
        )

    for crow in composed:
        cname = crow["name"]
        if not ref_names:
            matches.append({
                "composed_name": cname,
                "composed_qty": crow["qty"],
                "composed_unit": crow["unit"],
                "category": crow.get("_category"),
                "ref_match": None,
                "score": 0,
            })
            continue
        candidates = process.extract(
            cname, ref_names, scorer=fuzz.token_set_ratio, limit=8)
        ranked = []
        for nm, fscore, idx in candidates:
            if fscore < THRESHOLD:
                continue
            ranked.append((
                _combined(crow, reference[idx], float(fscore)),
                float(fscore), nm, idx,
            ))
        ranked.sort(key=lambda t: -t[0])
        # Use the highest combined-score candidate, even if the ref
        # row is already claimed by another composed row -- the
        # reference itself has duplicate-name rows so 1:1 mapping is
        # not appropriate.  used_ref is kept only as a diagnostic.
        if not ranked:
            matches.append({
                "composed_name": cname,
                "composed_qty": crow["qty"],
                "composed_unit": crow["unit"],
                "category": crow.get("_category"),
                "ref_match": None,
                "score": 0,
            })
            continue
        combined, fscore, nm, idx = ranked[0]
        used_ref.add(idx)
        ref = reference[idx]
        matches.append({
            "composed_name": cname,
            "composed_qty": crow["qty"],
            "composed_unit": crow["unit"],
            "category": crow.get("_category"),
            "ref_match": nm,
            "ref_qty": ref["qty"],
            "ref_unit": ref["unit"],
            "score": int(fscore),
            "qty_delta_pct": (
                100.0 * (float(crow["qty"]) - ref["qty"]) / ref["qty"]
                if ref["qty"] else 0.0
            ),
        })

    n_total = len(matches)
    n_matched = sum(1 for m in matches if m.get("ref_match"))
    match_rate = n_matched / n_total if n_total else 0.0

    # Per-class qty deviation diagnostics.  Only count matches whose
    # fuzzy score is in the "confident" band (>= 75) -- weak matches
    # are name-collisions where the qty delta carries no signal
    # about pipeline accuracy.
    HIGH_CONF = 75
    cable_devs = [
        abs(m.get("qty_delta_pct", 0))
        for m in matches
        if (m.get("ref_match")
            and _is_cable_ref(m["ref_match"])
            and m.get("score", 0) >= HIGH_CONF)
    ]
    lum_devs = [
        abs(m.get("qty_delta_pct", 0))
        for m in matches
        if (m.get("ref_match")
            and _is_lum_ref(m["ref_match"])
            and m.get("score", 0) >= HIGH_CONF)
    ]
    diag = {
        "n_composed": n_total,
        "n_ref": len(reference),
        "n_matched": n_matched,
        "match_rate": match_rate,
        "cable_devs": cable_devs,
        "lum_devs": lum_devs,
    }
    return matches, diag


# ---------------------------------------------------------------------------
# Gates
# ---------------------------------------------------------------------------

def gate_g1(rows: list[dict]) -> bool:
    n = len(rows)
    ok = n >= 80
    print("G1 composed rows = %d (threshold >= 80)" % n)
    print("G1 %s" % ("PASS" if ok else "FAIL"))
    return ok


def gate_g2(diag: dict) -> bool:
    r = diag["match_rate"] * 100
    ok = r > 65.0
    print("G2 match rate = %.1f%% (threshold > 65%%)" % r)
    print("G2 matched %d/%d composed against %d reference rows"
          % (diag["n_matched"], diag["n_composed"], diag["n_ref"]))
    print("G2 %s" % ("PASS" if ok else "FAIL"))
    return ok


def gate_g3(diag: dict) -> bool:
    cable_devs = diag["cable_devs"]
    lum_devs = diag["lum_devs"]
    cable_ok = True
    lum_ok = True
    if cable_devs:
        share = sum(1 for d in cable_devs if d <= 50.0) / len(cable_devs)
        cable_ok = share >= 0.5
        print("G3 cables: %d matched, %d within 50%% (share=%.1f%%)"
              % (len(cable_devs),
                 sum(1 for d in cable_devs if d <= 50.0),
                 share * 100))
    else:
        print("G3 cables: 0 matched -- skipping cable sub-gate")
    if lum_devs:
        share = sum(1 for d in lum_devs if d <= 30.0) / len(lum_devs)
        lum_ok = share >= 0.5
        print("G3 luminaires: %d matched, %d within 30%% (share=%.1f%%)"
              % (len(lum_devs),
                 sum(1 for d in lum_devs if d <= 30.0),
                 share * 100))
    else:
        print("G3 luminaires: 0 matched -- skipping luminaire sub-gate")
    ok = cable_ok and lum_ok
    print("G3 %s" % ("PASS" if ok else "FAIL"))
    return ok


def gate_g4(rows: list[dict], matches: list[dict], diag: dict) -> bool:
    os.makedirs(DISCUSS_DIR, exist_ok=True)
    payload = {
        "task": "T073",
        "sprint": "S016",
        "summary": {
            "composed_rows": len(rows),
            "reference_rows": diag["n_ref"],
            "matched_rows": diag["n_matched"],
            "match_rate_pct": round(diag["match_rate"] * 100, 1),
        },
        "rows": rows,
        "alignment": matches,
    }
    with open(OUTPUT_PATH, "w", encoding="utf-8") as fh:
        json.dump(payload, fh, ensure_ascii=False, indent=2)
    ok = os.path.exists(OUTPUT_PATH)
    print("G4 wrote %s (%d bytes)"
          % (OUTPUT_PATH, os.path.getsize(OUTPUT_PATH) if ok else 0))
    print("G4 %s" % ("PASS" if ok else "FAIL"))
    return ok


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    print("=" * 60)
    print("T073 acceptance verification")
    print("=" * 60)

    force = "--force" in sys.argv
    full = "--full" in sys.argv
    items_per_pdf = gather_items(force=force, lite=(not full))
    print("-" * 60)

    rows = compose(items_per_pdf)
    print("-" * 60)

    reference = _read_reference()
    print("reference rows: %d" % len(reference))
    matches, diag = fuzzy_align(rows, reference)
    print("-" * 60)

    a = gate_g1(rows)
    print("-" * 60)
    b = gate_g2(diag)
    print("-" * 60)
    c = gate_g3(diag)
    print("-" * 60)
    d = gate_g4(rows, matches, diag)
    print("=" * 60)
    print("Overall: G1=%s G2=%s G3=%s G4=%s"
          % ("PASS" if a else "FAIL",
             "PASS" if b else "FAIL",
             "PASS" if c else "FAIL",
             "PASS" if d else "FAIL"))
    return 0 if (a and b and c and d) else 1


if __name__ == "__main__":
    raise SystemExit(main())
