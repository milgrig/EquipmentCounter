"""
vor_docx_renderer.py — рендер ВОР в DOCX по эталонному формату ДБТ.

Эталон: «ВОР ЭО, Захватка 3_ГПК.docx»
Формат:
  • Альбомная A4, поля L=R=2см, T=0.8см, B=0.3см
  • Times New Roman 12pt в таблице, 14pt в преамбуле
  • Преамбула: стройка / объект / "Ведомость объемов работ" / основание / дата
  • Таблица 7 колонок: № п/п, Наименование вида работ, Ед.изм., РД,
    Формула расчета, Ссылка на чертежи, Дополнительная информация
  • Группировка строк по разделам (bold-строки):
    Щитовое оборудование / Светотехническое оборудование /
    Монтаж электроустановочных изделий / Кабельная продукция /
    Монтаж кабельных лотков / ПВХ изделия и трубы / Пусконаладочные работы
  • В конце: блок подписей "Составил _____ / Проверил _____"
"""

from __future__ import annotations

import io
import re
from datetime import date
from typing import Iterable

from docx import Document
from docx.shared import Cm, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ────────────────────────────────────────────────────────────────────────────
# Конфигурация формата
# ────────────────────────────────────────────────────────────────────────────

FONT_NAME = "Times New Roman"
TABLE_FONT_PT = 12
HEADER_FONT_PT = 14

# Ширины колонок (в см) — сумма ≈ 27.24 (ширина страницы A4 landscape - поля)
COL_WIDTHS_CM = [1.23, 9.76, 2.00, 1.50, 5.00, 4.25, 3.50]
COL_HEADERS = [
    "№\nп/п",
    "Наименование вида работ",
    "Ед.\nизм.",
    "Объем работ",
    "Формула расчета объемов работ и расхода материалов",
    "Ссылка на чертежи, спецификации",
    "Дополнительная информация",
]
COL_NUMBERS = ["1", "2", "3", "4", "5", "6", "7"]


# ────────────────────────────────────────────────────────────────────────────
# Группировка элементов ВОР по разделам
# ────────────────────────────────────────────────────────────────────────────

SECTION_RULES: list[tuple[str, list[re.Pattern]]] = [
    ("Щитовое оборудование", [
        re.compile(r"монтаж\s+щит", re.IGNORECASE),
        re.compile(r"подключени[ея]\s+жил", re.IGNORECASE),
        re.compile(r"щ[оа]?[оаб]?\b", re.IGNORECASE),
    ]),
    ("Светотехническое оборудование", [
        re.compile(r"светильник", re.IGNORECASE),
        re.compile(r"светодиодн", re.IGNORECASE),
        re.compile(r"монтаж\s+светил", re.IGNORECASE),
        re.compile(r"\bвыход\b", re.IGNORECASE),
        re.compile(r"световой\s+указат", re.IGNORECASE),
        re.compile(r"эвакуац", re.IGNORECASE),
    ]),
    ("Монтаж электроустановочных изделий", [
        re.compile(r"розетк", re.IGNORECASE),
        re.compile(r"выключател[ья]", re.IGNORECASE),
        re.compile(r"\bпост\s+управл", re.IGNORECASE),
        re.compile(r"датчик", re.IGNORECASE),
        re.compile(r"кнопк", re.IGNORECASE),
        re.compile(r"коробк", re.IGNORECASE),
    ]),
    ("Кабельная продукция", [
        re.compile(r"\bкабел[ья]\b", re.IGNORECASE),
        re.compile(r"провод", re.IGNORECASE),
        re.compile(r"\bввг\w*", re.IGNORECASE),
        re.compile(r"\bппг\w*", re.IGNORECASE),
        re.compile(r"\bвбшв\w*", re.IGNORECASE),
        re.compile(r"\bnym\b", re.IGNORECASE),
        re.compile(r"\bпвс\b", re.IGNORECASE),
    ]),
    ("Монтаж кабельных лотков и соединительных деталей", [
        re.compile(r"лот[ока]к?", re.IGNORECASE),
        re.compile(r"кабельн[аы]\w*\s+полк", re.IGNORECASE),
        re.compile(r"крышк[аи]\s+лотк", re.IGNORECASE),
    ]),
    ("ПВХ изделия и трубы", [
        re.compile(r"\bтруб[аы]\s+пвх", re.IGNORECASE),
        re.compile(r"\bгофр", re.IGNORECASE),
        re.compile(r"\bпнд\b", re.IGNORECASE),
        re.compile(r"\bпвх\b", re.IGNORECASE),
    ]),
    ("Пусконаладочные работы", [
        re.compile(r"измерени[ея]\s+сопротивл", re.IGNORECASE),
        re.compile(r"проверк[аи]\s+срабат", re.IGNORECASE),
        re.compile(r"целостност[иь]\s+жил", re.IGNORECASE),
        re.compile(r"\bпнр\b", re.IGNORECASE),
        re.compile(r"пусконаладоч", re.IGNORECASE),
        re.compile(r"\bлаборатор", re.IGNORECASE),
    ]),
]
SECTION_OTHER = "Прочие работы"

# ────────────────────────────────────────────────────────────────────────────
# T087 — Группировка светильников по высотным интервалам
# ────────────────────────────────────────────────────────────────────────────
# Канонический порядок и подписи для bucket-ов высот.
# Bucket-ключи совпадают со схемой height_bucketer.attribute_items
# (`<5m`, `5-13m`, `13-20m`, `20-35m`, `unknown`). Подписи берутся из
# эталонного ВОР ДБТ: «Монтаж светильников на шпильках к перекрытию
# на высоте до 5 метров» и т.д.
HEIGHT_BUCKET_ORDER: list[str] = ["<5m", "5-13m", "13-20m", "20-35m"]
HEIGHT_BUCKET_HEADER_FMT = (
    "Монтаж светильников на шпильках к перекрытию на высоте {phrase}"
)
HEIGHT_BUCKET_PHRASES: dict[str, str] = {
    "<5m": "до 5 метров",
    "5-13m": "от 5 до 13 метров",
    "13-20m": "от 13 до 20 метров",
    "20-35m": "от 20 до 35 метров",
}
HEIGHT_BUCKET_UNKNOWN = "unknown"

# Раздел, для которого включается под-группировка по высотам.
LUMINAIRE_SECTION_NAME = "Светотехническое оборудование"


def _elevation_to_bucket(elev_str: str) -> str:
    """Преобразовать «0.000» / «+4.200» / «+18.600» в bucket-ключ."""
    if not elev_str:
        return HEIGHT_BUCKET_UNKNOWN
    try:
        v = float(str(elev_str).replace(",", ".").lstrip("+"))
    except (ValueError, AttributeError):
        return HEIGHT_BUCKET_UNKNOWN
    if v < 5.0:
        return "<5m"
    if v < 13.0:
        return "5-13m"
    if v < 20.0:
        return "13-20m"
    if v < 35.0:
        return "20-35m"
    return HEIGHT_BUCKET_UNKNOWN


def _derive_height_bucket_from_elevations(per_elev: dict | None) -> str:
    """Доминирующий bucket для строки, у которой формула разложена по отметкам.

    per_elev: {"0.000": 14, "+4.200": 6, ...}. Если значения присутствуют в
    нескольких bucket-ах — выбираем bucket с максимальным суммарным total.
    """
    if not per_elev or not isinstance(per_elev, dict):
        return HEIGHT_BUCKET_UNKNOWN
    bucket_sums: dict[str, float] = {}
    for elev, qty in per_elev.items():
        if elev == "__other__":
            continue
        try:
            v = float(qty)
        except (TypeError, ValueError):
            continue
        if v <= 0:
            continue
        bkt = _elevation_to_bucket(elev)
        bucket_sums[bkt] = bucket_sums.get(bkt, 0.0) + v
    if not bucket_sums:
        return HEIGHT_BUCKET_UNKNOWN
    # max by sum; tie-break by canonical order (lowest bucket wins).
    best = max(
        bucket_sums.items(),
        key=lambda kv: (
            kv[1],
            -(HEIGHT_BUCKET_ORDER.index(kv[0])
              if kv[0] in HEIGHT_BUCKET_ORDER else 99),
        ),
    )
    return best[0]


def _resolve_height_bucket(item: dict) -> str:
    """Найти bucket для строки ВОР.

    Приоритеты (forward-compat с будущим vor_compose):
      1. item["_height_bucket"]      — если уже выставлено composer-ом.
      2. item["height_bucket"]       — если осталось от сырых элементов.
      3. derive_from_elevations(item["per_elevation"])  — из формулы по отметкам.
      4. "unknown".
    """
    for key in ("_height_bucket", "height_bucket"):
        v = item.get(key)
        if v and isinstance(v, str) and v.strip():
            return v.strip()
    return _derive_height_bucket_from_elevations(item.get("per_elevation"))


def _group_luminaires_by_height(
    rows: list[dict],
) -> list[tuple[str | None, list[dict]]]:
    """Разбить список строк раздела на под-группы по bucket-у высоты.

    Возвращает [(bucket_or_None, rows), ...] в каноническом порядке:
      <5m → 5-13m → 13-20m → 20-35m → unknown(None).
    Не светильники (строки, у которых имя не начинается со «Светильник»
    или «Монтаж светильник» и т.п.) идут в bucket=None (без под-заголовка).
    """
    lum_pat = re.compile(
        r"(?:^|\b)(?:монтаж\s+)?светил|светильник|световой\s+указат|"
        r"эвакуац|\bвыход\b",
        re.IGNORECASE,
    )
    buckets: dict[str, list[dict]] = {b: [] for b in HEIGHT_BUCKET_ORDER}
    unknown_lum: list[dict] = []
    non_lum: list[dict] = []
    for r in rows:
        nm = str(r.get("name") or "")
        if not lum_pat.search(nm):
            non_lum.append(r)
            continue
        b = _resolve_height_bucket(r)
        if b in buckets:
            buckets[b].append(r)
        else:
            unknown_lum.append(r)
    out: list[tuple[str | None, list[dict]]] = []
    for b in HEIGHT_BUCKET_ORDER:
        if buckets[b]:
            out.append((b, buckets[b]))
    if unknown_lum:
        out.append((HEIGHT_BUCKET_UNKNOWN, unknown_lum))
    if non_lum:
        out.append((None, non_lum))
    return out


def _classify_section(name: str) -> str:
    """Определить раздел ВОР по наименованию работы."""
    if not name:
        return SECTION_OTHER
    for sec_name, patterns in SECTION_RULES:
        for pat in patterns:
            if pat.search(name):
                return sec_name
    return SECTION_OTHER


def group_by_sections(items: list[dict]) -> list[tuple[str, list[dict]]]:
    """Сгруппировать строки агрегата по разделам в порядке SECTION_RULES."""
    buckets: dict[str, list[dict]] = {sec: [] for sec, _ in SECTION_RULES}
    buckets[SECTION_OTHER] = []
    for it in items:
        sec = _classify_section(it.get("name", ""))
        buckets[sec].append(it)
    out: list[tuple[str, list[dict]]] = []
    for sec, _ in SECTION_RULES:
        if buckets[sec]:
            out.append((sec, buckets[sec]))
    if buckets[SECTION_OTHER]:
        out.append((SECTION_OTHER, buckets[SECTION_OTHER]))
    return out


# ────────────────────────────────────────────────────────────────────────────
# Низкоуровневые помощники DOCX
# ────────────────────────────────────────────────────────────────────────────

def _set_cell_borders(cell):
    """Включить тонкие границы у ячейки."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tc_pr.append(borders)


def _set_cell_text(cell, text: str, *, bold=False, size_pt=TABLE_FONT_PT,
                   align=None, font=FONT_NAME):
    """Записать текст в ячейку с заданным форматом."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    # перенос строк
    parts = str(text or "").split("\n")
    for i, part in enumerate(parts):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(part)
        run.bold = bold
        run.font.name = font
        run.font.size = Pt(size_pt)
        # Cyrillic font hint
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), font)
        rFonts.set(qn('w:hAnsi'), font)
        rFonts.set(qn('w:cs'), font)


def _set_col_widths(table, widths_cm: list[float]):
    """Жёстко зафиксировать ширины колонок (cm)."""
    # tblLayout fixed
    tbl_pr = table._element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._element.insert(0, tbl_pr)
    layout = tbl_pr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')

    for row in table.rows:
        for ci, w in enumerate(widths_cm):
            if ci < len(row.cells):
                row.cells[ci].width = Cm(w)


# ────────────────────────────────────────────────────────────────────────────
# Главный рендер
# ────────────────────────────────────────────────────────────────────────────

def render_vor_docx(
    aggregated: list[dict],
    *,
    rel_folder: str = "",
    section_basis: str = "",
    drawing_prefix: str = "",
    issue_date: str | None = None,
    **_ignored,
) -> bytes:
    """Сформировать ВОР в формате эталона ДБТ. Возвращает байты .docx."""
    doc = Document()

    # 1. Section: A4 landscape, особые поля
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(0.8)
    sec.bottom_margin = Cm(0.3)

    if issue_date is None:
        issue_date = date.today().strftime("%d.%m.%Y")

    def _add_par(text, *, bold=False, align=None, size=HEADER_FONT_PT):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.name = FONT_NAME
        run.font.size = Pt(size)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:cs'), FONT_NAME)
        return p

    # 2. Преамбула (минимальная, по образцу ВОР_ЭО.docx)
    _add_par("Ведомость объемов работ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    if section_basis:
        _add_par(section_basis)
    _add_par(f"Дата составления {issue_date}г.")
    _add_par("")

    # 3. Таблица
    grouped = group_by_sections(aggregated)

    # T087: для раздела «Светотехническое оборудование» дополнительно
    # разбиваем строки на под-группы по высотным интервалам и добавляем
    # bold-параграф (строку-разделитель) перед каждой под-группой.
    # Поэтому plan-строки таблицы считаем с учётом sub-headers.
    luminaire_subgroups: dict[int, list[tuple[str | None, list[dict]]]] = {}
    total_section_rows = len(grouped)
    total_data_rows = 0
    total_sub_header_rows = 0
    for gi, (section_name, items) in enumerate(grouped):
        if section_name == LUMINAIRE_SECTION_NAME:
            subgroups = _group_luminaires_by_height(items)
            luminaire_subgroups[gi] = subgroups
            for bucket, rows in subgroups:
                # sub-header только для именованных bucket-ов
                if bucket in HEIGHT_BUCKET_PHRASES:
                    total_sub_header_rows += 1
                total_data_rows += len(rows)
        else:
            total_data_rows += len(items)

    # +1 шапка, +1 нумерация колонок
    total_rows = (
        2 + total_section_rows + total_sub_header_rows + total_data_rows
    )

    tbl = doc.add_table(rows=total_rows, cols=len(COL_HEADERS))
    tbl.style = "Table Grid"

    # шапка
    for ci, hdr in enumerate(COL_HEADERS):
        c = tbl.rows[0].cells[ci]
        _set_cell_text(c, hdr, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_borders(c)
    for ci, num in enumerate(COL_NUMBERS):
        c = tbl.rows[1].cells[ci]
        _set_cell_text(c, num, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_borders(c)

    # данные
    row_idx = 2
    item_no = 0  # сквозная нумерация по эталону

    def _write_section_row(text: str) -> None:
        """Записать строку-разделитель (bold во 2-й колонке)."""
        nonlocal row_idx
        for ci in range(len(COL_HEADERS)):
            cell = tbl.rows[row_idx].cells[ci]
            if ci == 1:
                _set_cell_text(cell, text, bold=True)
            else:
                _set_cell_text(cell, "")
            _set_cell_borders(cell)
        row_idx += 1

    def _write_data_row(it: dict) -> None:
        """Записать строку данных ВОР.

        Merge resolution (T089/S022_unblock):
          • outer structure preserved from T087 (S022) — factored helper
            so it can be called both from straight-section path and from
            T087 luminaire-by-height sub-grouping path.
          • inner row-format adopted from T078 (main):
              – strip "[UNMATCHED-LEGEND]" prefix from col 2 (deliverable
                must not show diagnostic markers; warnings belong to logs);
              – emit "" for col 5 (Formula) and col 7 (Additional info)
                to match T065_recon.md reference docx (130/130 rows empty).
        """
        nonlocal row_idx, item_no
        item_no += 1
        name = str(it.get("name", "")).strip()
        # T078 (main): strip "[UNMATCHED-LEGEND]" diagnostic prefix.
        name = _strip_unmatched_legend_prefix(name)
        unit = str(it.get("unit", "шт")).strip()
        total = it.get("total", 0)
        drawing_refs = str(it.get("drawing_refs", "")).strip()
        ref_text = _format_drawing_ref(drawing_refs, drawing_prefix)
        cells_text = [
            str(item_no),
            name,
            unit,
            _fmt_qty(total),
            "",            # T078: col 5 (Formula) empty per reference docx
            ref_text,
            "",            # T078: col 7 (Additional info) empty per reference docx
        ]
        for ci, text in enumerate(cells_text):
            cell = tbl.rows[row_idx].cells[ci]
            align = None
            if ci in (0, 2, 3):
                align = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_text(cell, text, align=align)
            _set_cell_borders(cell)
        row_idx += 1

    for gi, (section_name, items) in enumerate(grouped):
        # Section header row.
        _write_section_row(section_name)
        if gi in luminaire_subgroups:
            # T087: emit под-разделы по высоте.
            for bucket, rows in luminaire_subgroups[gi]:
                if bucket in HEIGHT_BUCKET_PHRASES:
                    phrase = HEIGHT_BUCKET_PHRASES[bucket]
                    sub_header = HEIGHT_BUCKET_HEADER_FMT.format(phrase=phrase)
                    _write_section_row(sub_header)
                for it in rows:
                    _write_data_row(it)
        else:
            for it in items:
                _write_data_row(it)

    _set_col_widths(tbl, COL_WIDTHS_CM)

    # 4. Подписи (удалены: данные исполнителей не известны автогенератору)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _fmt_qty(v) -> str:
    """Отформатировать количество (целые без .0, дробные с одной запятой)."""
    if v is None or v == "":
        return ""
    try:
        f = float(v)
    except (TypeError, ValueError):
        return str(v)
    if abs(f - round(f)) < 1e-6:
        return str(int(round(f)))
    return f"{f:.1f}".replace(".", ",")


# T078: pattern for the diagnostic warning prefix produced by web_app
# legend-coverage audit (web_app.py:1540).  We strip it before writing
# col 2 so the deliverable docx never contains the audit string.  Pattern
# is anchored at start and tolerates one or more whitespace chars after
# the closing bracket.
_UNMATCHED_LEGEND_RE = re.compile(r"^\s*\[UNMATCHED-LEGEND\]\s*", re.IGNORECASE)


def _strip_unmatched_legend_prefix(name: str) -> str:
    """Remove "[UNMATCHED-LEGEND] " prefix from a work-name if present.

    Returns the cleaned name; leaves any non-prefixed name unchanged.
    Idempotent over reruns (since the regex is anchored at start).
    """
    if not name:
        return name
    return _UNMATCHED_LEGEND_RE.sub("", name, count=1)


_LIST_NUM_RE = re.compile(r"л\.?\s*[\d.,\-–\s]+", re.IGNORECASE)


def _format_drawing_ref(raw: str, prefix: str) -> str:
    """Свернуть исходный список имён файлов в формат "<prefix>, л.<N>"."""
    if not raw:
        return ""
    # Если уже похож на эталон (содержит "л." и шифр) — оставить
    if "1Д-" in raw and "л." in raw:
        return raw
    # Извлечь номера листов из строк "005-План...", "006-План..."
    nums = []
    for tok in re.findall(r"\b(\d{2,3})\b", raw):
        if tok not in nums:
            nums.append(tok)
    # Также пустить через регулярку «л.NN»
    if not nums:
        return prefix
    # Сжать в диапазоны: 5,6,7,8 -> 5-8
    nums_int = sorted({int(n) for n in nums})
    ranges: list[str] = []
    i = 0
    while i < len(nums_int):
        j = i
        while j + 1 < len(nums_int) and nums_int[j + 1] == nums_int[j] + 1:
            j += 1
        if j > i:
            ranges.append(f"{nums_int[i]}-{nums_int[j]}")
        else:
            ranges.append(str(nums_int[i]))
        i = j + 1
    sheet_part = f"л.{','.join(ranges)}"
    if prefix:
        return f"{prefix}, {sheet_part}"
    return sheet_part
