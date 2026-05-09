#!/usr/bin/env python3
"""vor_compare_excel.py — Построчное сравнение эталонного ВОР и автоматически
сгенерированного ВОР с выводом в наглядный Excel-файл.

Формат таблиц ВОР (одинаковый для эталона и нашего):
    Колонки: № п/п | Наименование | Ед.изм. | Объём | Формула | Ссылка | Доп.инфо

Что делает скрипт:
    1. Читает обе .docx таблицы.
    2. Нормализует наименования (приводит синонимы, чистит регистр, убирает
       хвостовые номера позиций и т.п.).
    3. Делает fuzzy-сопоставление строк:
         - точное совпадение нормализованного имени → пара
         - иначе по similarity (difflib) с порогом 0.65
    4. Формирует Excel `vor_comparison.xlsx`:
         — Лист «Построчное сравнение»
             колонки: Раздел | Эталон: наименование / ед. / кол-во |
                      Наш:    наименование / ед. / кол-во |
                      Δ (абс) | Δ (%) | Статус | Заметка
         — Лист «Итого по разделам»
             precision/recall, сумма кол-в, средняя ошибка.
         — Лист «Только в эталоне» (пропущенные у нас)
         — Лист «Только у нас» (лишние позиции)

Запуск:
    python vor_compare_excel.py ЭТАЛОН.docx НАШ.docx -o vor_comparison.xlsx
"""

from __future__ import annotations

import argparse
import io
import re
import sys
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from pathlib import Path

# Ensure UTF-8 output on Windows console
if sys.platform == "win32":
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")
    except (AttributeError, ValueError):
        pass

try:
    from docx import Document
except ImportError:
    sys.exit("Нужен python-docx:  pip install python-docx")

try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    sys.exit("Нужен openpyxl:  pip install openpyxl")


# ---------------------------------------------------------------------------
#  Модель строки ВОР
# ---------------------------------------------------------------------------

@dataclass
class VorRow:
    idx: int                 # индекс строки в таблице (для отладки)
    section: str             # текущий раздел (подзаголовок выше по таблице)
    num: str                 # № п/п ("1", "2", ... либо "" для подпозиции)
    name: str                # исходное наименование
    name_norm: str           # нормализованное наименование
    unit: str
    qty: float | None
    qty_raw: str             # исходное представление ("35/1", "312", "" ...)
    ref: str = ""            # ссылка на чертежи
    extra: str = ""          # доп. информация
    is_section_header: bool = False

    def display(self, limit: int = 80) -> str:
        s = self.name.strip()
        return s if len(s) <= limit else s[: limit - 1] + "…"


# ---------------------------------------------------------------------------
#  Нормализация
# ---------------------------------------------------------------------------

_SYNONYMS: list[tuple[str, str]] = [
    # cable brands — разные написания
    (r"ВБШвнг\s*\(?А?\)?[-\s]*LS", "ВБШвнг(А)-LS"),
    (r"ВБШвнг\s*\(?А?\)?[-\s]*FRLS", "ВБШвнг(А)-FRLS"),
    (r"ППГнг\s*\(?A?А?\)?[-\s]*FRHF", "ППГнг(А)-FRHF"),
    (r"ППГнг\s*\(?A?А?\)?[-\s]*HF", "ППГнг(А)-HF"),
    (r"ВВГнг\s*\(?А?\)?[-\s]*LS", "ВВГнг(А)-LS"),
    # units
    (r"\bм\.?п\.?\b", "м"),
    (r"\bшт\.?\b", "шт"),
    (r"\bкомпл\.?\b", "комп"),
    # section separators
    (r"(\d)[xх×](\d)", r"\1х\2"),
    # whitespace normalization
    (r"\s+", " "),
]

_STOP_TOKENS = {
    "в", "на", "по", "с", "для", "и", "или", "до", "от", "при", "без",
    "шт", "м", "комп", "кг", "см", "мм", "каб", "изм",
}


def normalize_name(s: str) -> str:
    """Нормализовать наименование для fuzzy-матчинга."""
    if not s:
        return ""
    out = s.lower().strip()
    # убрать типичные шумы в начале: "1. ", "1) ", "№1 "
    out = re.sub(r"^\s*(?:№\s*)?\d+[\.\)]\s*", "", out)
    for pat, repl in _SYNONYMS:
        out = re.sub(pat, repl.lower(), out)
    # убрать пунктуацию (кроме цифр/букв/дефиса/плюса/запятой/точки)
    out = re.sub(r"[^\w\d\-\+,\.\sхx×/]", " ", out, flags=re.UNICODE)
    out = re.sub(r"\s+", " ", out).strip()
    return out


def _token_set(s: str) -> set[str]:
    toks = re.findall(r"[\w\d\-\+\.,/]+", s)
    return {t for t in toks if t not in _STOP_TOKENS and len(t) > 1}


def similarity(a: str, b: str) -> float:
    """Комбинированная мера: sequence ratio + jaccard по токенам."""
    if not a or not b:
        return 0.0
    sm = SequenceMatcher(None, a, b).ratio()
    ta, tb = _token_set(a), _token_set(b)
    if not ta or not tb:
        return sm
    jac = len(ta & tb) / len(ta | tb)
    return 0.5 * sm + 0.5 * jac


# ---------------------------------------------------------------------------
#  Чтение ВОР .docx
# ---------------------------------------------------------------------------

def _is_section_header(cells: list[str]) -> bool:
    """Раздел-подзаголовок: №, ед.изм., qty — все пустые, но есть наименование."""
    if not cells or len(cells) < 4:
        return False
    num, name, unit, qty = cells[0], cells[1], cells[2], cells[3]
    if not name:
        return False
    return (not num) and (not unit) and (not qty)


def _parse_qty(s: str) -> float | None:
    """Разобрать количество: '312' → 312; '35/1' → 35; '2,5' → 2.5."""
    if s is None:
        return None
    s = str(s).strip().replace(" ", "")
    if not s:
        return None
    s = s.replace(",", ".")
    if "/" in s:       # формат "35/1" (кол-во/типы) — берём первое
        s = s.split("/", 1)[0]
    m = re.match(r"[-+]?\d+(?:\.\d+)?", s)
    if not m:
        return None
    try:
        return float(m.group(0))
    except ValueError:
        return None


def read_vor_docx(path: Path) -> list[VorRow]:
    """Прочитать .docx ВОР и вернуть список строк."""
    doc = Document(str(path))
    if not doc.tables:
        raise ValueError(f"В файле {path.name} нет таблиц")

    # Берём самую большую таблицу (в ВОР обычно одна основная)
    table = max(doc.tables, key=lambda t: len(t.rows))

    rows: list[VorRow] = []
    current_section = ""

    # пропускаем шапку таблицы (1-2 первые строки: заголовки колонок и номера "1,2,3…")
    # определяем первую строку с данными эвристически
    start_idx = 0
    for i, row in enumerate(table.rows):
        first = row.cells[0].text.strip()
        second = row.cells[1].text.strip() if len(row.cells) > 1 else ""
        # шапка: "№" или первая текстовая строка "1" "2" "3"
        if "№" in first or "наименование" in second.lower():
            start_idx = i + 1
            continue
        # нумерная строка под шапкой ("1 | 2 | 3 | ...")
        if re.fullmatch(r"\d+", first) and re.fullmatch(r"\d+", second):
            start_idx = i + 1
            continue
        break

    for idx, row in enumerate(table.rows[start_idx:], start=start_idx):
        cells = [c.text.strip() for c in row.cells]
        # подравниваем до 7 колонок
        while len(cells) < 7:
            cells.append("")
        num, name, unit, qty_raw, formula, ref, extra = cells[:7]

        if not any(cells):
            continue

        if _is_section_header(cells):
            current_section = name
            rows.append(VorRow(
                idx=idx, section=current_section, num="", name=name,
                name_norm=normalize_name(name), unit="", qty=None,
                qty_raw="", ref=ref, extra=extra, is_section_header=True,
            ))
            continue

        rows.append(VorRow(
            idx=idx,
            section=current_section,
            num=num,
            name=name,
            name_norm=normalize_name(name),
            unit=unit,
            qty=_parse_qty(qty_raw),
            qty_raw=qty_raw,
            ref=ref,
            extra=extra,
            is_section_header=False,
        ))
    return rows


# ---------------------------------------------------------------------------
#  Сопоставление строк эталона и нашего файла
# ---------------------------------------------------------------------------

@dataclass
class Pair:
    """Результат сопоставления одной пары строк."""
    etalon: VorRow | None
    ours: VorRow | None
    score: float = 0.0
    status: str = ""          # "OK" | "DIFF_QTY" | "ONLY_ETALON" | "ONLY_OURS" | "SECTION"
    note: str = ""

    @property
    def section(self) -> str:
        if self.etalon:
            return self.etalon.section
        if self.ours:
            return self.ours.section
        return ""

    @property
    def delta_abs(self) -> float | None:
        if self.etalon and self.ours and self.etalon.qty is not None and self.ours.qty is not None:
            return self.ours.qty - self.etalon.qty
        return None

    @property
    def delta_pct(self) -> float | None:
        d = self.delta_abs
        if d is None or not self.etalon or not self.etalon.qty:
            return None
        return 100.0 * d / self.etalon.qty


_SIM_THRESHOLD = 0.60
_SECTION_BONUS = 0.10           # бонус за тот же раздел


def match_rows(etalon: list[VorRow], ours: list[VorRow]) -> list[Pair]:
    """Жадно сопоставить строки эталона со строками нашего файла.

    Стратегия:
      1. Секции-заголовки эталона сохраняем отдельной категорией (SECTION).
      2. Для каждой не-секции эталона ищем лучший match среди ours с similarity,
         учитывая совпадение раздела.
      3. Не сматченные строки ours уходят в ONLY_OURS.
    """
    pairs: list[Pair] = []
    used_ours: set[int] = set()

    # словарь индекс → строка нашего
    ours_data = [(i, r) for i, r in enumerate(ours) if not r.is_section_header]

    for er in etalon:
        if er.is_section_header:
            pairs.append(Pair(
                etalon=er, ours=None, score=1.0,
                status="SECTION", note="подзаголовок раздела",
            ))
            continue

        best_idx = -1
        best_score = 0.0
        for i, orow in ours_data:
            if i in used_ours:
                continue
            sc = similarity(er.name_norm, orow.name_norm)
            # бонус, если совпадает (нормализованный) раздел
            if er.section and orow.section and \
               normalize_name(er.section) == normalize_name(orow.section):
                sc += _SECTION_BONUS
            if sc > best_score:
                best_score = sc
                best_idx = i

        if best_idx >= 0 and best_score >= _SIM_THRESHOLD:
            used_ours.add(best_idx)
            orow = ours[best_idx]
            status, note = _assess_match(er, orow, best_score)
            pairs.append(Pair(
                etalon=er, ours=orow, score=round(best_score, 3),
                status=status, note=note,
            ))
        else:
            pairs.append(Pair(
                etalon=er, ours=None, score=round(best_score, 3),
                status="ONLY_ETALON",
                note=f"не найдено совпадение (лучший score={best_score:.2f})",
            ))

    # оставшиеся — лишние у нас
    for i, orow in ours_data:
        if i in used_ours:
            continue
        pairs.append(Pair(
            etalon=None, ours=orow, score=0.0,
            status="ONLY_OURS", note="нет в эталоне",
        ))

    return pairs


def _assess_match(er: VorRow, orow: VorRow, score: float) -> tuple[str, str]:
    """Определить статус пары по совпадению количества и единиц."""
    note_parts = []
    if er.unit and orow.unit and er.unit != orow.unit:
        note_parts.append(f"ед.изм: {er.unit!r}≠{orow.unit!r}")
    if er.qty is None and orow.qty is None:
        return ("OK", "; ".join(note_parts) or f"score={score:.2f}")
    if er.qty is None or orow.qty is None:
        note_parts.append("qty: одно из значений пустое")
        return ("DIFF_QTY", "; ".join(note_parts))
    if abs(er.qty - orow.qty) < 1e-6:
        return ("OK", "; ".join(note_parts) or "точное совпадение")
    # допуск 1%
    if er.qty and abs(orow.qty - er.qty) / max(abs(er.qty), 1e-9) < 0.01:
        return ("OK", "; ".join(note_parts) or "≤1%")
    return ("DIFF_QTY", "; ".join(note_parts) or "расхождение по qty")


# ---------------------------------------------------------------------------
#  Excel output
# ---------------------------------------------------------------------------

_STATUS_COLORS = {
    "OK":           "C6EFCE",   # зелёный
    "DIFF_QTY":     "FFEB9C",   # жёлтый
    "ONLY_ETALON":  "FFC7CE",   # красный
    "ONLY_OURS":    "BDD7EE",   # голубой
    "SECTION":      "D9D9D9",   # серый
}

_HEADER_FILL = PatternFill("solid", fgColor="305496")
_HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
_THIN = Side(style="thin", color="B4B4B4")
_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
_WRAP = Alignment(wrap_text=True, vertical="top")
_WRAP_CENTER = Alignment(wrap_text=True, vertical="top", horizontal="center")


def _write_header(ws, headers: list[str]):
    for i, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=i, value=h)
        c.fill = _HEADER_FILL
        c.font = _HEADER_FONT
        c.alignment = Alignment(wrap_text=True, vertical="center",
                                horizontal="center")
        c.border = _BORDER
    ws.row_dimensions[1].height = 32
    ws.freeze_panes = "A2"


def _auto_widths(ws, widths: list[int]):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w


def write_excel(
    pairs: list[Pair],
    etalon_path: Path,
    ours_path: Path,
    out_path: Path,
) -> None:
    wb = Workbook()

    # === Лист 1: Построчное сравнение ==========================
    ws = wb.active
    ws.title = "Построчное сравнение"
    headers = [
        "№",
        "Раздел",
        "Эталон: наименование",
        "Ед.",
        "Кол-во (эталон)",
        "Наш: наименование",
        "Ед.",
        "Кол-во (у нас)",
        "Δ абс.",
        "Δ %",
        "Score",
        "Статус",
        "Заметка",
    ]
    _write_header(ws, headers)
    _auto_widths(ws, [5, 22, 55, 8, 12, 55, 8, 12, 10, 10, 8, 14, 28])

    row_num = 2
    order_num = 0
    for p in pairs:
        order_num += 1
        section_name = p.section
        et_name = p.etalon.display(200) if p.etalon else ""
        et_unit = p.etalon.unit if p.etalon else ""
        et_qty = p.etalon.qty_raw if p.etalon else ""
        ou_name = p.ours.display(200) if p.ours else ""
        ou_unit = p.ours.unit if p.ours else ""
        ou_qty = p.ours.qty_raw if p.ours else ""
        d_abs = p.delta_abs
        d_pct = p.delta_pct

        values = [
            order_num,
            section_name,
            et_name,
            et_unit,
            et_qty,
            ou_name,
            ou_unit,
            ou_qty,
            "" if d_abs is None else round(d_abs, 2),
            "" if d_pct is None else f"{d_pct:+.1f}%",
            p.score if (p.etalon and p.ours) else "",
            p.status,
            p.note,
        ]
        fill = PatternFill("solid", fgColor=_STATUS_COLORS.get(p.status, "FFFFFF"))
        for i, v in enumerate(values, 1):
            c = ws.cell(row=row_num, column=i, value=v)
            c.fill = fill
            c.border = _BORDER
            c.alignment = _WRAP_CENTER if i in (1, 4, 5, 7, 8, 9, 10, 11, 12) else _WRAP

        # раздел-заголовок — выделяем жирным
        if p.status == "SECTION":
            for i in range(1, len(headers) + 1):
                cell = ws.cell(row=row_num, column=i)
                cell.font = Font(bold=True, italic=True)

        row_num += 1

    # === Лист 2: Итого по разделам =============================
    ws2 = wb.create_sheet("Итого по разделам")
    headers2 = [
        "Раздел",
        "Всего в эталоне",
        "Найдено у нас (OK)",
        "Расхождение кол-ва",
        "Пропущено",
        "Лишние у нас",
        "Recall, %",
        "Средняя |Δ%|",
    ]
    _write_header(ws2, headers2)
    _auto_widths(ws2, [42, 18, 20, 20, 14, 16, 12, 14])

    by_section: dict[str, dict[str, float]] = {}
    for p in pairs:
        if p.status == "SECTION":
            continue
        sec = p.section or "(без раздела)"
        agg = by_section.setdefault(sec, {
            "etalon_total": 0, "ok": 0, "diff": 0,
            "missed": 0, "extra": 0, "pct_sum": 0.0, "pct_cnt": 0,
        })
        if p.etalon:
            agg["etalon_total"] += 1
        if p.status == "OK":
            agg["ok"] += 1
        elif p.status == "DIFF_QTY":
            agg["diff"] += 1
            if p.delta_pct is not None:
                agg["pct_sum"] += abs(p.delta_pct)
                agg["pct_cnt"] += 1
        elif p.status == "ONLY_ETALON":
            agg["missed"] += 1
        elif p.status == "ONLY_OURS":
            agg["extra"] += 1

    r = 2
    totals = {"etalon_total": 0, "ok": 0, "diff": 0, "missed": 0,
              "extra": 0, "pct_sum": 0.0, "pct_cnt": 0}
    for sec, agg in by_section.items():
        total = agg["etalon_total"]
        recall = (agg["ok"] / total * 100.0) if total else 0
        avg_pct = (agg["pct_sum"] / agg["pct_cnt"]) if agg["pct_cnt"] else 0
        vals = [
            sec,
            int(total),
            int(agg["ok"]),
            int(agg["diff"]),
            int(agg["missed"]),
            int(agg["extra"]),
            f"{recall:.1f}%",
            f"{avg_pct:.1f}%" if agg["pct_cnt"] else "—",
        ]
        for i, v in enumerate(vals, 1):
            c = ws2.cell(row=r, column=i, value=v)
            c.border = _BORDER
            c.alignment = _WRAP_CENTER if i > 1 else _WRAP
        r += 1
        for k, v in agg.items():
            totals[k] += v

    # Total row
    total_total = totals["etalon_total"]
    total_recall = (totals["ok"] / total_total * 100.0) if total_total else 0
    total_avg_pct = (totals["pct_sum"] / totals["pct_cnt"]) if totals["pct_cnt"] else 0
    vals = [
        "ИТОГО",
        int(total_total),
        int(totals["ok"]),
        int(totals["diff"]),
        int(totals["missed"]),
        int(totals["extra"]),
        f"{total_recall:.1f}%",
        f"{total_avg_pct:.1f}%" if totals["pct_cnt"] else "—",
    ]
    for i, v in enumerate(vals, 1):
        c = ws2.cell(row=r, column=i, value=v)
        c.fill = PatternFill("solid", fgColor="305496")
        c.font = Font(bold=True, color="FFFFFF")
        c.border = _BORDER
        c.alignment = _WRAP_CENTER if i > 1 else _WRAP

    # === Лист 3: Только в эталоне =============================
    ws3 = wb.create_sheet("Только в эталоне")
    h3 = ["Раздел", "№", "Наименование", "Ед.", "Кол-во"]
    _write_header(ws3, h3)
    _auto_widths(ws3, [28, 6, 75, 8, 12])
    r = 2
    for p in pairs:
        if p.status != "ONLY_ETALON":
            continue
        e = p.etalon
        vals = [e.section, e.num, e.name, e.unit, e.qty_raw]
        fill = PatternFill("solid", fgColor=_STATUS_COLORS["ONLY_ETALON"])
        for i, v in enumerate(vals, 1):
            c = ws3.cell(row=r, column=i, value=v)
            c.fill = fill
            c.border = _BORDER
            c.alignment = _WRAP
        r += 1

    # === Лист 4: Только у нас =================================
    ws4 = wb.create_sheet("Только у нас")
    _write_header(ws4, h3)
    _auto_widths(ws4, [28, 6, 75, 8, 12])
    r = 2
    for p in pairs:
        if p.status != "ONLY_OURS":
            continue
        o = p.ours
        vals = [o.section, o.num, o.name, o.unit, o.qty_raw]
        fill = PatternFill("solid", fgColor=_STATUS_COLORS["ONLY_OURS"])
        for i, v in enumerate(vals, 1):
            c = ws4.cell(row=r, column=i, value=v)
            c.fill = fill
            c.border = _BORDER
            c.alignment = _WRAP
        r += 1

    # === Лист 5: Мета =========================================
    wsm = wb.create_sheet("Источники", 0)
    wsm["A1"] = "Файл"
    wsm["B1"] = "Путь"
    wsm["A1"].font = _HEADER_FONT
    wsm["A1"].fill = _HEADER_FILL
    wsm["B1"].font = _HEADER_FONT
    wsm["B1"].fill = _HEADER_FILL
    wsm["A2"] = "Эталон"
    wsm["B2"] = str(etalon_path)
    wsm["A3"] = "Наш"
    wsm["B3"] = str(ours_path)
    wsm.column_dimensions["A"].width = 14
    wsm.column_dimensions["B"].width = 120

    # Легенда статусов
    wsm["A5"] = "Легенда статусов"
    wsm["A5"].font = Font(bold=True)
    legend = [
        ("OK", "Строка найдена, количество совпадает (или Δ ≤ 1%)"),
        ("DIFF_QTY", "Строка найдена, но количество отличается"),
        ("ONLY_ETALON", "В эталоне есть, у нас пропущено"),
        ("ONLY_OURS", "У нас есть, в эталоне нет (лишнее)"),
        ("SECTION", "Подзаголовок раздела"),
    ]
    for i, (st, descr) in enumerate(legend, start=6):
        a = wsm.cell(row=i, column=1, value=st)
        b = wsm.cell(row=i, column=2, value=descr)
        a.fill = PatternFill("solid", fgColor=_STATUS_COLORS[st])
        a.border = _BORDER
        b.border = _BORDER

    wb.save(str(out_path))


# ---------------------------------------------------------------------------
#  CLI
# ---------------------------------------------------------------------------

def main():
    ap = argparse.ArgumentParser(
        description="Построчное сравнение эталонного и сгенерированного ВОР → Excel.",
    )
    ap.add_argument("etalon", help="Путь к эталонному .docx (человеческий ВОР)")
    ap.add_argument("ours", help="Путь к нашему сгенерированному .docx")
    ap.add_argument("-o", "--output", default="vor_comparison.xlsx",
                    help="Путь для выходного .xlsx (по умолчанию ./vor_comparison.xlsx)")
    args = ap.parse_args()

    etalon_path = Path(args.etalon)
    ours_path = Path(args.ours)
    out_path = Path(args.output)

    if not etalon_path.is_file():
        sys.exit(f"Не найден файл эталона: {etalon_path}")
    if not ours_path.is_file():
        sys.exit(f"Не найден наш файл: {ours_path}")

    print(f"Читаю эталон: {etalon_path.name}")
    etalon_rows = read_vor_docx(etalon_path)
    print(f"  строк: {len(etalon_rows)} "
          f"(из них разделов: {sum(1 for r in etalon_rows if r.is_section_header)})")

    print(f"Читаю наш:   {ours_path.name}")
    ours_rows = read_vor_docx(ours_path)
    print(f"  строк: {len(ours_rows)} "
          f"(из них разделов: {sum(1 for r in ours_rows if r.is_section_header)})")

    print("Сопоставляю строки…")
    pairs = match_rows(etalon_rows, ours_rows)

    stats = {}
    for p in pairs:
        stats[p.status] = stats.get(p.status, 0) + 1
    print("Итог сопоставления:")
    for k in ("OK", "DIFF_QTY", "ONLY_ETALON", "ONLY_OURS", "SECTION"):
        print(f"  {k:<12} {stats.get(k, 0)}")

    print(f"Пишу Excel: {out_path}")
    write_excel(pairs, etalon_path, ours_path, out_path)
    print("Готово.")


if __name__ == "__main__":
    main()
